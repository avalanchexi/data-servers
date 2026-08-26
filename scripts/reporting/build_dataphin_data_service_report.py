from __future__ import annotations

from pathlib import Path
from typing import Sequence

from PIL import Image, ImageDraw, ImageFont
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_volcengine_data_service_report as base


ROOT = Path(__file__).resolve().parents[2]
TMP = ROOT / ".tmp" / "dataphin-report"
IMAGES = TMP / "images"
OUTPUT = ROOT / "阿里云Dataphin数据服务应用架构与页面梳理.docx"
ARCH_DIAGRAM = TMP / "dataphin-application-architecture.png"
CALL_DIAGRAM = TMP / "dataphin-api-call-path.png"
OPS_CROP = IMAGES / "06-ops-monitoring-top.png"

OVERVIEW_SCREENSHOT = IMAGES / "01-overview-user-screenshot.png"
SCENARIO_FLOW = IMAGES / "02-official-scenario-flow.png"
API_MARKET = IMAGES / "03-api-market.png"
API_VERSION = IMAGES / "04-api-version-management.png"
API_WIZARD = IMAGES / "05-api-wizard-sort.png"
OPS_MONITORING = IMAGES / "06-ops-monitoring.png"
API_SQL = IMAGES / "07-direct-api-sql-editor.png"

BLUE = base.BLUE
DARK_BLUE = base.DARK_BLUE
INK = base.INK
MUTED = base.MUTED
LIGHT_BLUE = base.LIGHT_BLUE
LIGHT_GRAY = base.LIGHT_GRAY
CALLOUT = base.CALLOUT
WHITE = base.WHITE
GREEN = base.GREEN
GOLD = base.GOLD
RED = base.RED

SOURCES = [
    (
        "S1",
        "数据服务概述",
        "总览",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/summary-of-data-services",
        "定位、价值、优势、端到端使用流程",
    ),
    (
        "S2",
        "基本概念",
        "总览",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/the-basic-concept-1",
        "服务单元、应用、分组、API 服务、网络配置等定义",
    ),
    (
        "S3",
        "创建及管理服务项目",
        "服务管理",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/create-a-service-project",
        "项目隔离、成员、分组和发布管控",
    ),
    (
        "S4",
        "通过向导模式创建 API（服务单元）",
        "API 开发",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/create-an-api-service-unit-in-wizard-mode",
        "服务单元、请求/返回参数、排序和提交",
    ),
    (
        "S5",
        "通过直连数据源模式创建 API",
        "API 开发",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/create-an-api-by-using-the-direct-data-source-mode-operation-type",
        "SQL 模式、Basic/Dev-Prod、参数解析和试运行",
    ),
    (
        "S6",
        "管理 API",
        "API 开发",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/view-and-manage-apis",
        "版本、测试、发布、克隆、负责人和删除约束",
    ),
    (
        "S7",
        "数据服务市场",
        "API 市场",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/data-services-market",
        "已发布 API 的检索、文档、申请和下载",
    ),
    (
        "S8",
        "新建及管理我的应用",
        "应用管理",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/create-and-manage-service-applications",
        "应用主体、负责人、成员、AppKey/AppSecret",
    ),
    (
        "S9",
        "调用 API",
        "应用管理",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/call-api/",
        "API 文档、调试、开发/生产环境和调用示例",
    ),
    (
        "S10",
        "运维监控",
        "API 运维",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/operation-and-maintenance-monitoring/",
        "总览指标、异常影响、趋势、限流和告警",
    ),
    (
        "S11",
        "查看调用日志",
        "API 运维",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/view-call-log",
        "调用日志查询与故障排查",
    ),
    (
        "S12",
        "数据服务系统配置",
        "服务管理",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/system-configuration/",
        "鉴权、缓存、SQL 注入校验、日志与统计存储",
    ),
    (
        "S13",
        "Dataphin 支持的数据源",
        "数据来源",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/supported-data-sources",
        "数据服务可读取的数据源与环境映射",
    ),
    (
        "S14",
        "数据服务权限列表",
        "权限",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/data-service-permission-list-1",
        "管理员、开发、运维、应用角色的功能权限",
    ),
    (
        "S15",
        "API 资产详情",
        "资产治理",
        "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/api-asset-details",
        "API 类型、数据来源、调用模式、使用与变更信息",
    ),
]


def make_architecture_diagram() -> None:
    """Create an analysis diagram from the official functional relationships."""
    img = Image.new("RGB", (2400, 1600), "#" + WHITE)
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(base.FONT_CN_BOLD, 48)
    section_font = ImageFont.truetype(base.FONT_CN_BOLD, 29)
    body_font = ImageFont.truetype(base.FONT_CN, 24)
    small_font = ImageFont.truetype(base.FONT_CN, 21)

    draw.text(
        (100, 55),
        "Dataphin 数据服务应用架构（基于官方功能关系归纳）",
        font=title_font,
        fill="#" + INK,
    )

    base.rounded_box(
        draw,
        (80, 160, 430, 1390),
        "治理控制面",
        "服务项目\n成员与角色\n分组\n发布管控\n\n应用审批\nAPI 权限\n行级权限\n\n网络配置\n系统配置",
        "F7F9FB",
        DARK_BLUE,
        section_font,
        body_font,
    )

    base.rounded_box(
        draw,
        (540, 175, 1050, 445),
        "数据与服务生产",
        "业务数据源 / 物理表 / Dataphin 逻辑表\n服务单元（单表、多表、元数据配置）\n直连、服务单元、逻辑表、注册、组合 API",
        "F7FAFF",
        BLUE,
        section_font,
        small_font,
    )
    base.rounded_box(
        draw,
        (1140, 175, 1645, 445),
        "开发与发布",
        "向导模式 / SQL 模式\n参数解析、数据预览、调试与测试\n版本管理、发布校验、Dev-Prod 隔离",
        "F7FBF9",
        GREEN,
        section_font,
        small_font,
    )
    base.rounded_box(
        draw,
        (1735, 175, 2285, 445),
        "服务目录与消费",
        "API 市场：检索、文档、申请、下载\n应用：负责人、成员、AppKey/AppSecret\n企业应用、合作伙伴、开发者调用",
        "FFFBF3",
        GOLD,
        section_font,
        small_font,
    )
    base.arrow(draw, (1050, 310), (1140, 310), GREEN)
    base.arrow(draw, (1645, 310), (1735, 310), GOLD)
    base.arrow(draw, (430, 310), (540, 310), DARK_BLUE)

    base.rounded_box(
        draw,
        (540, 590, 1035, 905),
        "接入与网关",
        "公网 / VPC 域名\nAPI 网关或内置网关\n白名单、token 鉴权、AppKey/AppSecret\n同步 / 异步调用、JSON 返回",
        "F7FAFF",
        BLUE,
        section_font,
        small_font,
    )
    base.rounded_box(
        draw,
        (1130, 590, 1655, 905),
        "运行时控制",
        "路由与版本\n鉴权、限流、超时\n缓存、SQL 注入校验\n行级权限与数据范围控制",
        "FFF8F8",
        RED,
        section_font,
        small_font,
    )
    base.rounded_box(
        draw,
        (1750, 590, 2285, 905),
        "数据执行与返回",
        "开发环境读取开发数据源（Dev-Prod）\n生产环境读取生产数据源\n执行 SQL / 已注册服务\n结果、错误码和调用标识返回",
        "F7FBF9",
        GREEN,
        section_font,
        small_font,
    )
    base.arrow(draw, (1035, 750), (1130, 750), RED)
    base.arrow(draw, (1655, 750), (1750, 750), GREEN)
    draw.line((2010, 445, 2010, 510, 790, 510), fill="#" + BLUE, width=6)
    base.arrow(draw, (790, 510), (790, 590), BLUE, width=6)

    base.rounded_box(
        draw,
        (540, 1050, 2285, 1380),
        "可观测性与运营闭环",
        "运维总览（发布、在线、调用、错误、offline、受影响应用）\n异常影响 Top 10、访问趋势、API/应用限流、告警规则\n调用日志：请求、响应、SQL、耗时、错误信息（受网关及日志采集配置约束）",
        "FFF8F8",
        RED,
        section_font,
        body_font,
    )
    base.arrow(draw, (1400, 905), (1400, 1050), RED)
    base.arrow(draw, (430, 1220), (540, 1220), DARK_BLUE)

    draw.text(
        (90, 1490),
        "读图：上半部分是服务生产与消费控制面；中部是线上调用执行面；底部将运行数据反馈到限流、告警、版本和权限治理。",
        font=small_font,
        fill="#" + MUTED,
    )
    img.save(ARCH_DIAGRAM, quality=95)


def make_call_diagram() -> None:
    img = Image.new("RGB", (2400, 720), "#" + WHITE)
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(base.FONT_CN_BOLD, 42)
    box_title = ImageFont.truetype(base.FONT_CN_BOLD, 26)
    box_body = ImageFont.truetype(base.FONT_CN, 21)
    draw.text((90, 45), "API 运行时调用链（架构归纳）", font=title_font, fill="#" + INK)

    boxes = [
        ((70, 170, 410, 520), "业务应用", "应用身份\nAppKey / AppSecret\n请求参数"),
        ((500, 170, 840, 520), "网络入口", "公网 / VPC\n域名与白名单\nAPI 网关"),
        ((930, 170, 1270, 520), "安全与流控", "token 鉴权\n权限 / 行级权限\n限流 / 超时 / 缓存"),
        ((1360, 170, 1700, 520), "API 运行时", "版本路由\n同步 / 异步\nSQL 或注册服务"),
        ((1790, 170, 2130, 520), "数据与响应", "开发/生产数据源\nJSON / 错误码\n调用标识"),
    ]
    colors = [GOLD, BLUE, RED, GREEN, DARK_BLUE]
    for (box, title, body), color in zip(boxes, colors):
        base.rounded_box(draw, box, title, body, "F7F9FB", color, box_title, box_body)
    for idx in range(len(boxes) - 1):
        x1 = boxes[idx][0][2]
        y = (boxes[idx][0][1] + boxes[idx][0][3]) // 2
        x2 = boxes[idx + 1][0][0]
        base.arrow(draw, (x1, y), (x2, y), colors[idx + 1], width=6)
    draw.line((1100, 540, 1100, 620, 2050, 620, 2050, 540), fill="#" + RED, width=5)
    draw.text(
        (1170, 625),
        "调用统计与明细日志回流到 API 运维，用于异常分析、限流和告警",
        font=box_body,
        fill="#" + RED,
    )
    img.save(CALL_DIAGRAM, quality=95)


def crop_ops_image() -> None:
    with Image.open(OPS_MONITORING) as img:
        width, height = img.size
        bottom = min(height, 1750)
        img.crop((0, 0, width, bottom)).save(OPS_CROP)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def setup_document():
    doc, bullet_id, decimal_id = base.setup_document()
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    clear_paragraph(header)
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    run = header.add_run("阿里云 Dataphin | 数据服务应用架构")
    base.set_font(run, bold=True, color=INK, size=8.8)
    run = header.add_run("\t官方文档梳理")
    base.set_font(run, color=MUTED, size=8.8)

    footer = section.footer.paragraphs[0]
    clear_paragraph(footer)
    run = footer.add_run("资料访问日期：2026-08-25")
    base.set_font(run, size=9, color=MUTED)
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    base.add_page_number(footer)
    return doc, bullet_id, decimal_id


def add_source_paragraph(doc, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("来源：")
    base.set_font(run, size=8.5, color=MUTED)
    base.add_hyperlink(p, label, url, color=BLUE, size=8.5)


def add_figure(doc, path: Path, caption: str, alt: str, *, width: float = 6.35,
               source_label: str | None = None, source_url: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    base.set_picture_alt(p, alt)
    base.add_caption(doc, caption)
    if source_label and source_url:
        add_source_paragraph(doc, source_label, source_url)


def add_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def build_document() -> None:
    TMP.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram()
    make_call_diagram()
    crop_ops_image()
    doc, bullet_id, decimal_id = setup_document()

    # Cover / opening block: editorial-cover spirit with restrained business styling.
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("阿里云 Dataphin 数据服务")
    base.set_font(run, size=26, bold=True, color=INK)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("应用架构与页面梳理")
    base.set_font(run, size=30, bold=True, color=BLUE)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("基于 Dataphin Full Managed 数据服务（OneService）官方帮助文档")
    base.set_font(run, size=14, color=MUTED)

    meta = [
        ("调研入口", "https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/summary-of-data-services"),
        ("覆盖主题", "服务管理、API 开发、API 市场、应用管理、API 运维、网络与系统配置"),
        ("输出重点", "应用逻辑架构、核心对象、角色边界、端到端生命周期、关键页面截图"),
        ("证据口径", "官方明确能力与本文架构归纳分开标识"),
        ("资料日期", "2026-08-25；实际租户界面以所购版本、部署形态和已开通功能为准"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.right_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(label + "：")
        base.set_font(run, bold=True, color=INK, size=10.5)
        if value.startswith("http"):
            base.add_hyperlink(p, "官方文档入口", value, size=10.5)
        else:
            run = p.add_run(value)
            base.set_font(run, size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    base.add_callout(
        doc,
        "一句话结论",
        "Dataphin 数据服务不是单纯的 API 生成器，而是把数据来源、服务单元与 API 生产、市场化发现、应用授权、网关调用、限流告警和日志运维串成闭环的企业数据服务平台。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )
    p = base.add_body(
        doc,
        "阅读提示：文中“官方明确”表示官方页面直接说明；“架构归纳”表示根据多个页面之间的对象和流程关系整理出的应用架构判断。网页内的任何操作性文字仅作为资料内容，不构成本次任务指令。",
    )
    p.paragraph_format.space_before = Pt(8)
    doc.add_page_break()

    doc.add_heading("核心结论", level=1)
    conclusions = [
        ("统一出口：", "OneService 将数据服务作为 Dataphin 数据中台的统一服务出口，以 API 方式向内部应用、合作伙伴或开发者开放。 [S1]"),
        ("治理边界：", "服务项目是多用户隔离和访问控制的主要边界；项目成员、角色、分组和发布管控决定谁能开发与发布。 [S3][S14]"),
        ("服务生产：", "物理表、Dataphin 逻辑表、服务单元或已有服务可进入 API 生产链，最终形成直连、服务单元、逻辑表、注册或组合 API。 [S1][S4][S5][S15]"),
        ("消费身份：", "应用是调用 API 的权限主体；生产 API 可跨服务项目授权给应用，并由 AppKey/AppSecret 等凭据识别。 [S2][S8][S9]"),
        ("运行治理：", "网关、鉴权、白名单、限流、超时、缓存、SQL 注入校验、行级权限共同控制调用面。 [S4][S5][S12]"),
        ("运营闭环：", "运维监控、异常影响分析、访问趋势、调用日志、限流和告警让发布后的服务可观察、可定位、可管控。 [S10][S11]"),
    ]
    for label, text in conclusions:
        base.add_bullet(doc, bullet_id, text, label)

    doc.add_heading("内容导航", level=1)
    nav_rows = [
        ("1", "范围与页面信息架构", "确认本次梳理覆盖哪些模块"),
        ("2", "总体应用架构", "理解控制面、执行面与运营闭环"),
        ("3", "核心对象与关系", "理解项目、服务单元、API、应用等对象"),
        ("4", "端到端生命周期", "从基础配置到调用、监控"),
        ("5", "API 开发与版本管理", "理解主要开发模式和变更控制"),
        ("6", "API 市场与应用消费", "理解发现、申请、授权与调用"),
        ("7", "安全、运维与治理", "理解运行态的保护和可观测性"),
        ("8", "角色与职责", "理解各角色在平台中的操作边界"),
        ("9", "落地建议与注意事项", "将官方能力映射到企业实施"),
    ]
    base.add_table(
        doc,
        ["章节", "主题", "阅读目的"],
        nav_rows,
        [900, 3000, 5460],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    doc.add_heading("1 文档范围与页面信息架构", level=1)
    base.add_body(
        doc,
        "本报告以用户提供的“数据服务概述”页面为入口，沿左侧目录和顶部“服务”菜单梳理数据服务相关能力。页面目录按功能模块组织，本报告则按“生产—发布—发现—授权—调用—运维”的业务价值流重新组合。",
    )
    add_figure(
        doc,
        OVERVIEW_SCREENSHOT,
        "图 1  用户提供的阿里云官方文档入口截图：左侧为数据服务目录",
        "阿里云 Dataphin 数据服务概述帮助文档截图，左侧目录包含数据服务概述、基本概念、服务管理、API 开发、API 市场与应用管理、应用运维和 API 运维。",
        source_label="数据服务概述",
        source_url="https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/summary-of-data-services",
    )
    module_rows = [
        ("服务管理", "项目、成员、分组、应用、网络、系统配置", "建立组织、权限、网络和全局运行边界", "治理控制面"),
        ("API 开发", "服务单元、API、元数据、版本、测试与发布", "把数据或已有服务加工成可发布接口", "服务生产面"),
        ("API 市场", "已发布 API、检索、文档、权限申请、下载", "让服务可发现、可比较、可申请", "服务目录面"),
        ("应用管理", "我的应用、成员、密钥、已授权 API、调用说明", "形成可识别的消费主体和凭据", "消费控制面"),
        ("应用运维", "应用调用情况与日志", "从应用视角查看消费行为", "消费运营面"),
        ("API 运维", "运维监控、限流、告警、调用日志", "监控 API 健康度并处置异常", "平台运营面"),
    ]
    base.add_table(
        doc,
        ["页面域", "主要对象/页面", "回答的问题", "架构定位"],
        module_rows,
        [1500, 3000, 3060, 1800],
    )

    doc.add_heading("2 总体应用架构", level=1)
    add_figure(
        doc,
        ARCH_DIAGRAM,
        "图 2  Dataphin 数据服务总体应用架构（架构归纳）",
        "Dataphin 数据服务应用架构图，包括治理控制面、数据与服务生产、开发与发布、服务目录与消费、接入网关、运行时控制、数据执行与返回、可观测性与运营闭环。",
    )
    base.add_callout(
        doc,
        "架构归纳",
        "官方文档以功能页面描述产品能力，并未发布一张名为“应用架构”的统一拓扑图。图 2 是根据项目、服务单元、API、应用、网关和运维对象之间的关系进行的逻辑分层，不能替代具体租户的部署拓扑。",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("2.1 六层逻辑架构", level=2)
    layer_rows = [
        ("治理层", "服务项目、成员、角色、分组、发布管控", "定义所有权、隔离和职责分离", "控制面"),
        ("数据与服务资产层", "物理表、逻辑表、服务单元、注册服务", "形成可被 API 使用的数据与服务契约", "生产面"),
        ("API 产品层", "直连、服务单元、逻辑表、注册、组合 API", "将查询或既有服务封装成标准接口", "生产面"),
        ("目录与消费层", "API 市场、应用、授权、API 文档", "发现服务并把调用权授予具体应用", "消费控制面"),
        ("接入与运行层", "公网/VPC、网关、鉴权、限流、缓存、执行", "安全稳定地处理请求并返回结果", "执行面"),
        ("运维运营层", "指标、趋势、日志、告警、限流、影响分析", "把运行数据反馈到治理和变更决策", "闭环"),
    ]
    base.add_table(doc, ["层次", "核心构件", "主要职责", "平面"], layer_rows, [1800, 2800, 3160, 1600])

    doc.add_heading("2.2 控制面与执行面", level=2)
    base.add_bullet(doc, bullet_id, "服务项目、角色、发布管控、API 市场、权限申请、应用与系统配置决定“谁可以生产、发布、发现和调用什么”。", "控制面：")
    base.add_bullet(doc, bullet_id, "业务应用通过公网或 VPC 访问域名，经网关鉴权、限流和路由进入 API 运行时，再访问对应环境的数据源或已注册服务。", "执行面：")
    base.add_bullet(doc, bullet_id, "调用指标和明细日志进入运维页面，用于异常影响分析、趋势、限流和告警，并反向驱动发布、版本或权限治理。", "反馈闭环：")
    base.add_body(doc, "“控制面/执行面/反馈闭环”是本文的架构归纳；其组成能力均可在官方页面找到，但官方没有统一使用这三个名称。")

    doc.add_heading("3 核心对象与关系", level=1)
    object_rows = [
        ("服务项目", "成员、角色、服务单元、API、分组", "多用户隔离和访问控制的主要边界", "服务管理"),
        ("服务单元", "一个或多个物理表及统一字段元数据", "为向导/SQL API 提供稳定字段集合", "API 开发"),
        ("元数据", "非结构化表的二维结构化配置", "让特定数据源可被服务单元统一使用", "API 开发"),
        ("API", "请求/返回参数、数据来源、版本、状态、运行环境", "可发布、授权和运维的服务产品", "API 开发"),
        ("应用", "负责人、成员、AppKey/AppSecret、已授权 API", "调用生产 API 的权限主体", "应用管理"),
        ("授权关系", "API ↔ 应用", "控制应用能否调用以及权限期限", "市场/权限"),
        ("网络配置", "域名、VPC 白名单、网关环境", "调用成功的基础网络条件", "服务管理"),
        ("调用记录", "请求、响应、SQL、耗时、错误、应用", "故障排查和运营分析的证据", "API 运维"),
    ]
    base.add_table(doc, ["对象", "包含/关联", "业务意义", "主要页面"], object_rows, [1500, 3000, 3160, 1700])
    base.add_callout(
        doc,
        "关键关系",
        "服务项目 1:N 服务单元与 API；服务单元 N:1 或 N:N 关联物理表；API N:N 授权给应用；应用使用密钥经网络和网关调用；每次调用产生统计或明细日志。服务项目是生产侧边界，应用是消费侧边界。",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("4 端到端生命周期", level=1)
    add_figure(
        doc,
        SCENARIO_FLOW,
        "图 3  官方场景与使用流程：按角色分工完成基础配置、开发、调用和监控",
        "Dataphin 数据服务官方场景流程图，按超级管理员、开发用户、业务应用用户和运维用户展示创建服务项目、配置网络、创建服务单元、开发与发布 API、创建应用、申请与调用 API、监控 API 的流程。",
        source_label="数据服务概述中的场景及使用流程",
        source_url="https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/summary-of-data-services",
    )
    lifecycle = [
        ("基础治理：", "超级管理员或系统管理员创建服务项目、添加成员并配置网络；项目发布管控决定高风险 API 变更是否阻塞。"),
        ("数据建模：", "按数据源、物理表或 Dataphin 逻辑表准备数据；必要时先建立单表/多表服务单元和字段元数据。"),
        ("服务开发：", "选择服务单元向导、SQL、直连数据源、逻辑表、注册或组合方式创建 API，并配置参数、超时、调用模式、版本等。"),
        ("测试发布：", "提交后在开发/测试链路验证；Dev-Prod 模式下开发读取开发数据源、发布后读取生产数据源。"),
        ("市场发现：", "发布后的 API 进入数据服务市场，用户可检索、查看生产文档、申请权限或下载 API 文档。"),
        ("应用授权：", "创建应用并由审批形成消费主体；应用负责人维护成员和密钥，API 权限与应用绑定。"),
        ("在线调用：", "应用经公网/VPC 和网关，使用 AppKey/AppSecret 或相应 token 调用生产 API；开发环境可切换版本调试。"),
        ("运维闭环：", "运维人员查看调用量、错误率、offline、异常影响应用和趋势，结合调用日志、限流和告警处理问题。"),
    ]
    for label, text in lifecycle:
        base.add_numbered(doc, decimal_id, text, label)

    doc.add_heading("5 API 开发与版本管理", level=1)
    doc.add_heading("5.1 API 类型与生产方式", level=2)
    api_rows = [
        ("直连数据源 API", "直接编写 SQL 访问数据源", "查询，及特定增删改能力", "快速封装已有数据库逻辑"),
        ("服务单元 API", "基于服务单元，支持向导或 SQL", "单表/多表字段、查询条件、排序", "复用统一字段契约"),
        ("逻辑表 API", "基于 Dataphin 逻辑表", "标准逻辑模型与环境映射", "承接数仓标准化数据"),
        ("注册 API", "将已有后端服务注册到平台", "转发调用、超时、返回路径", "纳管既有服务"),
        ("组合 API", "引用多个子 API 形成组合逻辑", "复杂服务组装", "提高服务复用度"),
    ]
    base.add_table(doc, ["类型", "数据/服务来源", "主要能力", "适用目的"], api_rows, [1700, 2500, 2860, 2300])
    base.add_body(doc, "API 资产详情页明确列出直连数据源、服务单元、逻辑表、注册和组合五类 API；具体可用类型和操作能力可能受版本及增值功能影响。 [S15]")

    doc.add_heading("5.2 向导模式：用字段配置生成查询接口", level=2)
    add_figure(
        doc,
        API_WIZARD,
        "图 4  服务单元向导模式：选择字段，配置请求参数、返回参数和排序",
        "Dataphin 服务单元 API 向导页面，左侧显示服务单元字段，右侧配置请求参数、返回参数和排序设置。",
        source_label="通过向导模式创建 API（服务单元）",
        source_url="https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/create-an-api-service-unit-in-wizard-mode",
    )
    base.add_bullet(doc, bullet_id, "向导模式适合标准查询；字段从同一服务单元选择，并配置对外参数名、绑定字段、类型、操作符、示例和是否必填。", "官方明确：")
    base.add_bullet(doc, bullet_id, "分页查询应设置稳定排序，避免结果重复或丢失；请求中的 OrderByList 与页面排序字段会共同影响顺序。", "设计关注：")

    doc.add_heading("5.3 SQL 模式：用脚本承载复杂逻辑", level=2)
    add_figure(
        doc,
        API_SQL,
        "图 5  直连数据源 SQL 编辑与试运行结果页面",
        "Dataphin 直连数据源 API 参数配置页面，包含 Basic 模式、数据源、SQL 编辑区、解析参数、SQL 试运行以及 JSON 返回结果。",
        source_label="通过直连数据源模式创建 API",
        source_url="https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/create-an-api-by-using-the-direct-data-source-mode-operation-type",
    )
    base.add_bullet(doc, bullet_id, "SQL 模式可以解析请求和返回参数，支持基础 SQL 与高级 SQL；高级 SQL 使用 MyBatis 风格标签承载动态条件。", "官方明确：")
    base.add_bullet(doc, bullet_id, "Basic 模式在开发、提交和发布阶段均读取生产数据；Dev-Prod 模式开发/提交读取开发数据、发布后读取生产数据。", "环境边界：")
    base.add_bullet(doc, bullet_id, "Create/Update/Delete 等操作类 API 会占用更多计算资源，需结合事务、批次、并行度和最大输入条数控制风险。", "资源边界：")

    doc.add_heading("5.4 版本、测试与发布", level=2)
    add_figure(
        doc,
        API_VERSION,
        "图 6  API 版本管理：同一 API 下并存已发布与已提交版本",
        "Dataphin API 版本管理面板，显示 API 名称、版本号、版本状态、创建人、更新时间以及查看、测试、对比、编辑、删除等操作。",
        source_label="管理 API",
        source_url="https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/view-and-manage-apis",
    )
    base.add_bullet(doc, bullet_id, "已发布版本不可直接编辑；平台可基于线上版本创建新版本，支持测试、版本对比、克隆、发布与历史版本管理。", "版本模型：")
    base.add_bullet(doc, bullet_id, "新版本发布会校验新增必填请求参数、删减请求/返回参数、类型或调用类型变化，并按项目发布管控决定阻塞或放行。", "变更门禁：")
    base.add_callout(doc, "架构判断", "API_ID 稳定而配置版本可演进，使下游应用不必因每次数据源或查询逻辑调整而修改对接地址；但兼容性变更仍需发布管控和影响通知。", fill=LIGHT_BLUE)

    doc.add_heading("6 API 市场与应用消费", level=1)
    doc.add_heading("6.1 市场：服务发现和权限申请入口", level=2)
    add_figure(
        doc,
        API_MARKET,
        "图 7  API 市场：搜索区、项目分类、筛选功能区和 API 列表",
        "Dataphin API 市场页面，标注搜索区、数据服务项目分类、筛选与功能区、API 列表，可查看 API 文档和申请权限。",
        source_label="数据服务市场",
        source_url="https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/data-services-market",
    )
    base.add_bullet(doc, bullet_id, "市场只呈现已发布 API，并提供搜索、按服务项目浏览、筛选、查看生产 API 文档、立即申请和批量下载文档。", "官方明确：")
    base.add_bullet(doc, bullet_id, "API 市场将“开发资产”转成“可发现的服务产品”，是连接生产者和消费者的服务目录。", "架构归纳：")

    doc.add_heading("6.2 应用：消费身份和密钥边界", level=2)
    application_rows = [
        ("应用对象", "名称、分组、场景、负责人、成员", "确定调用责任主体和维护人"),
        ("凭据", "AppKey、AppSecret，可自动生成或按网关条件自定义", "识别应用并完成签名/token 生成"),
        ("API 权限", "生产 API 可跨服务项目授权给应用", "把服务消费权落到具体应用"),
        ("有效期", "应用成员和权限均可配置期限", "降低长期闲置权限风险"),
        ("密钥变更", "应用负责人可重置 AppSecret", "需要同步调用方，避免 SDK 中断"),
    ]
    base.add_table(doc, ["构件", "官方能力", "架构作用"], application_rows, [1800, 3900, 3660])
    base.add_callout(doc, "消费边界", "服务项目解决“谁能生产和管理 API”，应用解决“哪个业务系统可以调用 API”。两者通过 API 授权连接，形成生产侧和消费侧相互独立的责任边界。", fill=LIGHT_BLUE)

    doc.add_heading("6.3 运行时调用链", level=2)
    add_figure(
        doc,
        CALL_DIAGRAM,
        "图 8  API 运行时调用链：应用身份、网关控制、API 执行与日志回流（架构归纳）",
        "Dataphin API 调用链架构图，从业务应用通过公网或 VPC 和 API 网关，经 token 鉴权、权限、限流、超时和缓存，到 API 运行时与数据源，结果和调用日志返回。",
    )
    base.add_body(doc, "调用时，应用先准备请求参数和凭据，经域名与网关进入平台；运行时按环境和版本路由到相应数据源或注册服务，返回 JSON/错误信息，并将统计或明细日志送入 API 运维。 [S9][S11][S12]")

    doc.add_heading("7 安全、运维与治理", level=1)
    doc.add_heading("7.1 运行时保护", level=2)
    control_rows = [
        ("身份鉴权", "AppKey/AppSecret、SDK/token；部分网关可配置 token 鉴权", "阻止匿名调用", "与网关环境相关"),
        ("网络边界", "公网/VPC 域名、VPC 白名单", "限制可达网络", "公共云用户需关注白名单"),
        ("权限控制", "API 申请授权、应用成员、项目角色、行级权限", "限制可调用 API 和可返回数据", "需正确审批与回收"),
        ("流量与超时", "API/应用限流、同步/异步、服务超时", "保护平台和后端数据源", "操作类 API 资源成本更高"),
        ("缓存", "系统 Redis+内存、仅内存、指定 Redis", "降低后端查询压力", "缓存配置与容量需匹配"),
        ("SQL 安全", "SQL 注入校验；SQL/参数解析约束", "降低恶意输入风险", "注册 API 不支持该校验"),
        ("发布管控", "兼容性校验、应用授权/组合引用阻塞策略", "避免版本变更破坏下游", "策略按项目配置"),
    ]
    base.add_table(doc, ["控制点", "能力", "目的", "边界"], control_rows, [1500, 2860, 2700, 2300])

    doc.add_heading("7.2 运维监控页面", level=2)
    add_figure(
        doc,
        OPS_CROP,
        "图 9  API 运维监控页面（顶部区域）：核心指标、异常影响和访问趋势",
        "Dataphin API 运维监控页面顶部截图，包含发布 API 数、在线 API 数、调用 API 数、在线调用率、调用次数、异常次数、错误率、offline 百分比、异常影响应用和调用趋势图。",
        source_label="运维监控",
        source_url="https://help.aliyun.com/zh/dataphin/fullmanaged/user-guide/operation-and-maintenance-monitoring/",
    )
    ops_points = [
        ("总览指标：", "发布 API 数、在线 API 数、调用 API 数、在线 API 调用率、调用次数、异常次数、错误率、offline 百分比、异常 API 和受影响应用。"),
        ("影响分析：", "异常影响应用 Top 10 和调用异常次数 API Top 10，帮助按业务影响优先级排查。"),
        ("趋势分析：", "将调用量、异常量、异常 API 数和受影响应用数按时间关联，判断异常范围和演化。"),
        ("处置动作：", "从 API 或应用维度配置限流与告警，并下钻调用日志查看错误码、请求响应、SQL 和耗时。"),
    ]
    for label, text in ops_points:
        base.add_bullet(doc, bullet_id, text, label)
    base.add_callout(
        doc,
        "日志完整性边界",
        "调用统计和明细日志是否完整取决于网关类型、是否开通日志采集、历史 API 是否重新发布、日志存储位置和保留时长。官方明确提示部分场景可能出现日志缺失或少量冗余，不能把“页面无日志”直接等同于“未发生调用”。 [S11][S12]",
        fill=CALLOUT,
        accent=RED,
    )

    doc.add_heading("8 角色与职责", level=1)
    role_rows = [
        ("超级管理员", "全局", "管理所有 API/项目调用情况及平台级能力", "全局治理与审计"),
        ("系统管理员", "租户/服务管理", "创建服务项目、配置网络与系统参数", "基础平台配置"),
        ("服务项目管理员", "项目", "成员、分组、API/服务单元管理、发布策略、运维", "项目责任人"),
        ("开发用户", "项目", "创建服务单元、元数据和 API，测试与发布", "服务生产者"),
        ("运维用户", "负责项目", "查看生产调用情况、日志、限流和告警", "运行保障"),
        ("应用负责人", "应用", "应用、成员、密钥和已授权 API 的维护", "消费侧责任人"),
        ("应用普通成员", "应用", "使用已授权 API、查看应用调用分析和日志", "服务消费者"),
        ("业务/合作伙伴开发者", "外部或内部系统", "基于 API 文档和 SDK 发起调用", "下游集成者"),
    ]
    base.add_table(doc, ["角色", "作用域", "主要职责", "架构定位"], role_rows, [1600, 1700, 3760, 2300])
    base.add_callout(doc, "职责分离", "平台将项目管理、API 开发、API 运维和应用维护分配给不同角色。企业落地时应至少区分服务生产者、发布/运维责任人和消费应用负责人，并为密钥、审批与告警建立明确交接。", fill=LIGHT_BLUE)

    doc.add_heading("9 落地建议与注意事项", level=1)
    recommendations = [
        ("先划项目边界：", "按业务域、团队责任和数据访问范围划分服务项目，再设置成员和发布管控；不要用一个大项目承载所有数据服务。"),
        ("把服务单元当作契约：", "统一字段名、类型、说明、行级权限与最大返回条数，避免每个 API 重复解释同一张表。"),
        ("按复杂度选择生产方式：", "标准查询优先向导，复杂查询使用 SQL；已有接口采用注册，跨 API 逻辑采用组合，避免所有场景都落到自由 SQL。"),
        ("将应用作为授权最小单元：", "一个业务系统或清晰责任域对应一个应用，避免多个系统共用 AppKey/AppSecret；配置有效期并定期轮换。"),
        ("建立兼容性门禁：", "生产 API 被应用授权或被组合 API 引用后，默认使用阻塞式发布管控；对必填参数、返回字段和类型变化做下游验证。"),
        ("日志与告警先于上线：", "在大规模发布前确认网关日志采集、存储位置、保留期、告警渠道和责任人，否则出现故障时难以还原调用现场。"),
        ("区分 Basic 与 Dev-Prod：", "Basic 开发阶段也可能读取生产数据；需要测试隔离时应使用 Dev-Prod，并验证开发/生产数据源映射。"),
        ("核对版本与增值功能：", "数据服务、行级权限、增删改操作、网关和日志能力存在购买、版本或部署限制，实施前应按当前租户控制台复核。"),
    ]
    for label, text in recommendations:
        base.add_numbered(doc, decimal_id, text, label)

    doc.add_heading("9.1 重要产品边界", level=2)
    boundary_rows = [
        ("数据源功能变化", "概述页提示 Dataphin 数据源功能已下线，可使用 Dataphin JDBC Driver 替代", "不要把旧截图中的“Dataphin 数据源”入口当作当前必然可用功能"),
        ("界面版本差异", "官方截图来自不同页面更新时间和产品版本", "菜单名、字段和样式可能与当前租户不同"),
        ("部署拓扑未知", "官方帮助页描述功能，不公开具体租户底层组件部署", "图 2 仅是逻辑应用架构，不推断实例、集群或网络拓扑"),
        ("日志并非绝对完整", "网关模式、SLS/日志采集、历史发布状态会影响统计与明细", "排障时结合网关、后端和应用侧日志交叉验证"),
        ("权限和增值开关", "数据服务、行级权限、操作类 API 等受购买与权限条件限制", "先做功能可用性清单，再设计流程"),
    ]
    base.add_table(doc, ["边界", "官方依据", "对架构梳理的影响"], boundary_rows, [1800, 3600, 3960])

    doc.add_heading("10 架构总结", level=1)
    base.add_callout(
        doc,
        "最终判断",
        "Dataphin OneService 的核心是建立一条受治理的数据服务供应链：服务项目规定生产边界，服务单元和多种 API 模式提高生产效率，版本与发布管控保护变更，API 市场完成发现，应用与密钥形成消费身份，网关和运行时控制保障访问，监控与日志形成运营闭环。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )
    base.add_body(doc, "如果映射到企业应用架构，可将其理解为“数据服务开发平台 + API 产品目录 + 应用授权中心 + 网关运行时 + 服务运营平台”的组合。其价值不只在于把 SQL 变成 HTTP 接口，更在于把接口的所有权、权限、版本、消费主体和运行责任放在同一生命周期内。")

    doc.add_heading("附录 A 主要官方来源", level=1)
    base.add_body(doc, "以下页面用于本报告的事实核对与页面截图，均为阿里云官方帮助文档，访问日期为 2026-08-25。")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    base.set_table_geometry(table, [650, 2200, 1700, 4810])
    for idx, title in enumerate(["编号", "官方页面", "模块", "用于本报告的证据"]):
        base.shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        base.set_cell_text(table.rows[0].cells[idx], title, bold=True, color=INK, size=9.3)
    base.set_repeat_table_header(table.rows[0])
    for sid, title, module, url, purpose in SOURCES:
        row = table.add_row()
        base.set_cell_text(row.cells[0], sid, align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[1].text = ""
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        base.add_hyperlink(p, title, url, size=9.2)
        base.set_cell_text(row.cells[2], module)
        base.set_cell_text(row.cells[3], purpose)

    doc.add_heading("附录 B 截图清单与证据边界", level=1)
    screenshot_rows = [
        ("图 1", "用户提供的帮助文档入口截图", "确认目录和页面范围", "原图保留"),
        ("图 3", "数据服务概述：场景及使用流程", "确认角色分工与主流程", "阿里云官方静态图片"),
        ("图 4", "向导模式 API 参数页面", "确认服务单元字段、入参、出参和排序", "阿里云官方静态图片"),
        ("图 5", "直连数据源 SQL 编辑页面", "确认 SQL、参数解析和试运行", "阿里云官方静态图片"),
        ("图 6", "API 版本管理", "确认版本状态与操作", "阿里云官方静态图片"),
        ("图 7", "API 市场", "确认搜索、项目、筛选和列表", "阿里云官方静态图片"),
        ("图 9", "运维监控", "确认指标、异常影响与趋势", "官方原图裁取顶部关键区域"),
    ]
    base.add_table(doc, ["图号", "页面", "用于说明", "处理方式"], screenshot_rows, [900, 2700, 3560, 2200])

    boundary_rows = [
        ("官方明确", "对象定义、角色权限、页面入口、配置项、使用限制、操作流程和指标", "可在附录 A 对应页面直接找到"),
        ("架构归纳", "六层逻辑架构、控制面/执行面、调用链、核心关系和落地建议", "由多个官方页面之间的关系推导"),
        ("未覆盖", "具体租户的部署拓扑、组织现状、实际 QPS/SLA、成本和已购功能清单", "需结合企业控制台与现状另行调研"),
    ]
    base.add_table(doc, ["口径", "内容", "证据说明"], boundary_rows, [1500, 5000, 2860])

    doc.core_properties.title = "阿里云 Dataphin 数据服务应用架构与页面梳理"
    doc.core_properties.subject = "基于阿里云 Dataphin 官方文档的数据服务应用架构分析"
    doc.core_properties.keywords = "阿里云, Dataphin, OneService, 数据服务, 应用架构, API, 服务单元"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUTPUT)


if __name__ == "__main__":
    required: Sequence[Path] = (
        OVERVIEW_SCREENSHOT,
        SCENARIO_FLOW,
        API_MARKET,
        API_VERSION,
        API_WIZARD,
        OPS_MONITORING,
        API_SQL,
    )
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required images:\n" + "\n".join(missing))
    build_document()
    print(OUTPUT)
