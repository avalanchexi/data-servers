from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "华为云DataArts Studio数据服务应用架构梳理（含官网截图）.docx"
EXPECTED_HEADINGS = {
    "执行摘要",
    "1 总体应用架构",
    "2 部署与网络架构",
    "3 核心对象模型",
    "4 API提供方：服务生产与生命周期",
    "5 API调用方：服务发现、授权与安全",
    "6 治理、可观测与运营架构",
    "7 开放接口与集成架构",
    "8 非功能边界与选型判断",
    "9 建议的企业落地蓝图",
    "10 结论",
    "附录A  官网“数据服务”页面地图",
    "附录B  主要官方来源",
}


def twips(value) -> int:
    return round(value.twips)


def fail(message: str) -> None:
    raise AssertionError(message)


def main() -> None:
    if not REPORT.exists() or REPORT.stat().st_size < 100_000:
        fail("Report is missing or unexpectedly small")

    doc = Document(REPORT)
    if len(doc.sections) != 1:
        fail(f"Expected one section, got {len(doc.sections)}")
    section = doc.sections[0]
    dimensions = {
        "page_width": (twips(section.page_width), 12240),
        "page_height": (twips(section.page_height), 15840),
        "top_margin": (twips(section.top_margin), 1440),
        "bottom_margin": (twips(section.bottom_margin), 1440),
        "left_margin": (twips(section.left_margin), 1440),
        "right_margin": (twips(section.right_margin), 1440),
    }
    for label, (actual, expected) in dimensions.items():
        if abs(actual - expected) > 2:
            fail(f"{label}: expected {expected} twips, got {actual}")

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_text = [cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
    body_text = "\n".join(paragraphs + table_text)
    missing = sorted(EXPECTED_HEADINGS - set(paragraphs))
    if missing:
        fail(f"Missing headings: {missing}")
    if len(body_text) < 11_000:
        fail(f"Document text is unexpectedly short: {len(body_text)} characters")
    if re.search(r"\b(?:TODO|TBD|FIXME)\b|\{\{.+?\}\}|<placeholder>", body_text, re.I):
        fail("Placeholder content remains in the document")

    if len(doc.inline_shapes) != 17:
        fail(f"Expected 17 embedded figures, got {len(doc.inline_shapes)}")

    root = doc._element
    doc_pr_nodes = root.xpath(".//wp:docPr")
    missing_alt = [node.get("name", "unnamed") for node in doc_pr_nodes if not node.get("descr", "").strip()]
    if missing_alt:
        fail(f"Figures missing alternative text: {missing_alt}")

    hyperlink_nodes = root.xpath(".//w:hyperlink")
    if len(hyperlink_nodes) < 20:
        fail(f"Expected at least 20 hyperlinks, got {len(hyperlink_nodes)}")

    if len(doc.tables) < 15:
        fail(f"Expected at least 15 tables, got {len(doc.tables)}")
    for index, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_w is None or tbl_w.get(qn("w:w")) != "9360":
            fail(f"Table {index} does not use 9360 dxa width")
        if tbl_ind is None or tbl_ind.get(qn("w:w")) != "120":
            fail(f"Table {index} does not use 120 dxa indent")
        grid = table._tbl.tblGrid
        grid_total = sum(int(col.get(qn("w:w"))) for col in grid.gridCol_lst)
        if grid_total != 9360:
            fail(f"Table {index} grid totals {grid_total}, expected 9360")
        for row_index, row in enumerate(table.rows, start=1):
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                fail(f"Table {index} row {row_index} can split across pages")
            if tr_pr.find(qn("w:trHeight")) is not None:
                fail(f"Table {index} row {row_index} has fixed height")

    page_map_tables = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers[:3] == ["序号", "层级", "页面标题"]:
            page_map_tables.append(table)
    if len(page_map_tables) != 1:
        fail(f"Expected one page-map table, got {len(page_map_tables)}")
    page_ids = [int(row.cells[0].text.strip()) for row in page_map_tables[0].rows[1:]]
    if page_ids != list(range(1, 53)):
        fail("The official page map does not contain the exact sequence 1..52")

    with zipfile.ZipFile(REPORT) as archive:
        names = archive.namelist()
        media = [name for name in names if name.startswith("word/media/")]
        if len(media) != 17:
            fail(f"Expected 17 media parts, got {len(media)}")
        numbering = archive.read("word/numbering.xml").decode("utf-8")
        if "w:abstractNum" not in numbering or "w:num" not in numbering:
            fail("Word numbering definitions are missing")

    print(f"PASS report={REPORT}")
    print(f"size_bytes={REPORT.stat().st_size}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"figures={len(doc.inline_shapes)}")
    print(f"hyperlinks={len(hyperlink_nodes)}")
    print(f"official_page_map_rows={len(page_ids)}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"FAIL {exc}", file=sys.stderr)
        raise
