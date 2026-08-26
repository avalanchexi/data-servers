from __future__ import annotations

import json
from collections import OrderedDict
from pathlib import Path

from PIL import Image
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_volcengine_data_service_report import (
    BLUE,
    CALLOUT,
    DARK_BLUE,
    GOLD,
    GREEN,
    INK,
    LIGHT_BLUE,
    MUTED,
    RED,
    SCREENSHOT,
    SOURCES,
    add_body,
    add_bullet,
    add_callout,
    add_caption,
    add_hyperlink,
    add_table,
    hex_rgb,
    set_font,
    set_picture_alt,
    setup_document,
)


ROOT = Path(__file__).resolve().parents[2]
TMP = ROOT / ".tmp" / "volcengine-data-service-report"
CONTENT_JSON = TMP / "page-content" / "page-content.json"
INVENTORY_JSON = TMP / "page-screenshots" / "inventory.json"
IMAGE_DIR = TMP / "page-screenshots"
DIAGRAM = TMP / "data-service-business-architecture.png"
OUTPUT = ROOT / "火山引擎DataLeap数据服务业务架构梳理（官网层级·全功能版）.docx"


PAGE_OVERVIEWS = {
    "S1": "说明数据服务要解决的业务问题、平台价值和产品优势，形成服务创建、管理、运维与共享的总体定位。",
    "S2": "定义业务线、项目、物理表、逻辑表和应用等核心对象，明确组织边界、资产映射与授权主体。",
    "S3": "以“数据源接入—物理表—逻辑表—API 创建/开发/测试/发布—应用授权—调用”为完整上手链路。",
    "S4": "注册数据库连接并进行连通性测试，覆盖多类 MySQL、ByteHouse、Doris 与 StarRocks 数据源。",
    "S5": "在数据库负载或异常场景下，将指定应用/API 的流量按比例迁移到目标数据源。",
    "S6": "把源库表元数据注册为平台物理表，并管理字段、安全等级、标签、负责人、版本与删除。",
    "S7": "以物理表映射生成逻辑表，屏蔽存储差异，并提供字段治理、查询约束、授权、探查、主备和版本管理。",
    "S8": "覆盖 API 文件夹、新建、三种开发模式、参数、高级配置、保存、测试、发布、版本、灰度和详情。",
    "S9": "把 API、条件、函数与合并节点组合为串并行工作流，并完成测试、发布、版本与运维配置。",
    "S10": "说明 API 发布后的环境地址、密钥与授权、请求头/请求体、调用示例、返回结构、字段类型和错误类型。",
    "S11": "覆盖已发布 API 的限流、授权回收、血缘、报警规则与调用监控。",
    "S12": "定义 OneService 占位符及日期/分区函数，并说明适用限制、参数类型和输出规则。",
    "S13": "定义 Dynamic SQL 元素、表达式字面量、运算符、内置函数、闭包与切片能力。",
    "S14": "提供已上架 API 的搜索、详情、跨项目调用申请、行列权限、授权、测试和运维入口。",
    "S15": "提供同业务线逻辑表的搜索、详情、探查和跨项目权限申请，减少重复建设。",
    "S16": "定义租户、业务线、项目三级角色权限，并支持权限添加、申请、审批、撤销和删除。",
    "S17": "以项目作为隔离与访问控制边界，管理创建申请、审批策略、成员、应用授权和封禁策略。",
    "S18": "管理业务域分组及业务线管理员，为同业务线资源跨项目复用建立上层边界。",
    "S19": "管理 API 的应用授权主体及静态/OAuth 2.0 密钥，覆盖应用与密钥的增改查删。",
    "S20": "绑定辅助网卡、私网与公网 IP，使已发布 API 可以通过公网访问。",
    "S21": "配置允许调用 API 的 VPC、子网、域名和资源组，仅支持同 VPC 网络调用。",
    "S22": "集中处理我的审核与我的申请，覆盖筛选、详情、通过、驳回和撤销。",
    "S23": "以公开/私有标签组对物理表、逻辑表和 API 分类，支持标签维护与项目关联。",
}


PAGE_FUNCTIONS = {
    "S1": [
        "识别异构存储接口、权限与高可用、服务运维、同质接口重复建设四类问题",
        "提供向导式、脚本式和原生 SQL 三种服务生产方式",
        "提供 API 版本、鉴权、资产转移、报警和调用监控",
        "统一接口标准并收敛数据计算逻辑、计算资源与存储资源",
        "通过审核、鉴权和流控降低敏感数据暴露与合规成本",
        "以“配置即服务”自动生产和部署数据服务",
    ],
    "S2": [
        "业务线：同类业务的组织集合，可包含多个项目并支持经审批的跨项目复用",
        "项目：带业务语义的协作与隔离单元，承载数据源和 API 的规范化管理",
        "物理表：注册到平台的源表元数据，是查询 DSL 构造所依赖的底层对象",
        "逻辑表：物理表的虚拟映射，承担字段转换、规范命名和容灾映射",
        "应用：API 授权与调用的主体，由应用管理统一维护",
    ],
    "S3": [
        "准备业务线、项目、角色权限和公网/VPC 网络",
        "注册数据源并打通网络连接",
        "创建物理表并注册库表/字段元数据",
        "创建逻辑表并完成逻辑建模",
        "选择脚本式、向导式或原生式创建 API",
        "配置逻辑表、请求/返回参数和缓存等高级参数",
        "测试查询逻辑并查看结果/日志",
        "发布到测试或线上环境并授权应用",
        "从 API 详情获取地址、代码示例和接口文档并发起调用",
    ],
    "S4": [
        "支持 Cloud/VEDB MySQL、RDS MySQL、ByteHouse、ByteHouse CE、EMR-Doris、EMR-StarRocks、EMR 全托管 StarRocks",
        "支持 VPC 或公网连接，并按数据源类型配置白名单、安全组和端口",
        "录入数据源名称、描述、实例/集群、数据库、账号和连接池参数",
        "执行连通性测试并注册数据源",
        "按名称、负责人和类型搜索数据源",
        "编辑 Owner 与连接参数",
        "发起数据源切流",
        "删除不再使用的数据源",
    ],
    "S5": [
        "从源数据源选择目标数据源",
        "按调用应用筛选受影响范围",
        "识别已上线 API 及“API 需修改”情况",
        "设置 0–100 的切流比例",
        "处理多数据源 Join 与逻辑表/物理表映射风险",
        "关闭切流后再修改目标数据源",
    ],
    "S6": [
        "从已注册数据源选择库表并批量加入已选列表",
        "维护字段安全等级、主键和实时更新字段",
        "设置表安全等级，并可同步创建同名逻辑表",
        "查看基础信息、字段明细和变更记录",
        "刷新源表字段并比较差异",
        "回退到历史字段版本",
        "按名称、ID、负责人、数据源类型和标签搜索",
        "维护标签、负责人并删除物理表",
    ],
    "S7": [
        "使用极简模式批量创建逻辑表或标准模式创建单表",
        "选择物理表并维护逻辑字段、中文名、描述和安全等级",
        "设置查询约束、默认授权项目、目录、描述和下游 API 发布管控",
        "查看基础信息、字段、抽样数据、SQL 探查、授权和变更记录",
        "编辑字段名称、类型、主键、指标/维度关联等属性",
        "在不影响下游 API 的前提下切换来源物理表并重新映射字段",
        "编辑默认授权、安全等级、查询约束、标签和发布管控",
        "配置同类型、同字段结构的整表主备关系并切换线上表",
        "在无关联 API 时删除逻辑表",
        "搜索、定位、移动、变更负责人和管理目录树",
    ],
    "S8": [
        "创建最多三级 API 文件夹并支持重命名、移动和删除",
        "新建脚本式、向导式或原生式 API",
        "配置 API 名称、目录、最大 QPS、安全等级、最大 limit、报警模板和初始版本",
        "脚本式支持同源最多 10 张逻辑表、普通 SQL 和 Dynamic SQL",
        "脚本式支持请求/返回参数解析、手工校正、排序、示例值、缺省值和安全等级",
        "向导式支持便携式数据源直连或标准式逻辑表，并自动生成查询语句",
        "原生式支持同源最多 5 张逻辑表和调用方动态 SQL",
        "高级配置支持 JSON/JSONCompact、结果缓存和向导式分页",
        "支持直接保存、另存新版本和另存新 API",
        "测试支持 dryRun、表单/JSON 请求、数组/结构参数、结果/日志/请求体/状态查看",
        "支持将测试参数同步到 API 详情调用信息",
        "发布到测试环境或按审批流程发布到线上环境",
        "编辑 API 基本信息与标签",
        "查看版本、发布环境、下线、删除、重新发布和版本对比",
        "按版本、应用或 IP 配置灰度比例与启停状态",
        "从 API 详情进入调用、授权、报警和监控",
    ],
    "S9": [
        "创建 API 编排目录和编排对象",
        "配置名称、目录、负责人、安全等级和描述",
        "画板默认包含开始与结束节点，并支持拖拽添加节点",
        "支持节点命名/描述、画布缩放、定位和自动布局",
        "支持添加/删除依赖、在线路中插入节点和删除节点",
        "支持开始、结束、函数、分支、API、合并六类节点",
        "结束节点定义编排整体返回参数，并可由测试结果解析同步",
        "函数节点调用 API 网关触发器并可独立测试",
        "分支节点依据输入参数配置条件走向",
        "API 节点复用已开发 API",
        "合并节点支持 append 或 merge，并处理来源字段与别名",
        "以开始节点参数测试整个工作流并同步结果参数",
        "发布到测试/线上环境并维护基本配置与历史版本",
        "编排详情复用调用、授权和报警能力；当前不支持灰度策略",
    ],
    "S10": [
        "查看各环境调用地址、路径、请求代码和生成接口文档",
        "创建应用密钥并向 API 授权应用",
        "在 VPC 或公网环境发起 POST 调用",
        "配置 user、env、URL、ApiID、Params/Sql 等调用参数",
        "按 API 配置行级和列级权限、最大 QPS 与有效期",
        "查看 Fields、Data、TimeCost、DataCnt、ErrorCode、RenderedSql 和 BaseResp",
        "查看 FieldType 类型枚举和 QueryErrorType 错误类型",
    ],
    "S11": [
        "设置 API 最大 QPS 并启用限流",
        "新增应用授权并配置有效期、最大 QPS、调用负责人、优先级和行列权限",
        "编辑授权或回收调用权限",
        "查看逻辑表—API—应用—底层物理表/存储的血缘关系",
        "新增报警规则并配置应用、等级、周期、指标、接收人、Webhook、短信和状态",
        "导入报警模板并维护默认失败率报警规则",
        "编辑或删除报警规则",
        "按应用、环境、时间范围和时区查看 QPS、成功率与 PCT99",
    ],
    "S12": [
        "支持 ${} 直接替换和 #{} 按类型生成 SQL 片段两类占位符",
        "#{} 用于降低 SQL 注入风险并支持数组参数",
        "TO_YYYYMMDD：把日期转为 YYYYmmdd",
        "ADD_DATE/SUB_DATE：按偏移量调整日期",
        "DATE_PARTITION：根据逻辑表就绪时间生成分区日期",
        "DATE_SUB_PARTITION：按二级分区条件生成就绪日期",
        "函数可结合两类占位符使用，但只计算用户输入参数",
        "不支持分号收尾、窗口函数及部分 MySQL With/grouping sets/lateral view/Union 语法",
    ],
    "S13": [
        "Dynamic SQL 最外层使用 <select>，并处理 XML 特殊字符转义",
        "支持 if 条件片段",
        "支持 choose/when/otherwise 单分支选择",
        "支持 trim/where 自动处理前后缀、AND/OR 和分隔符",
        "支持 foreach 遍历数组、集合或 Map 并生成 IN 等片段",
        "支持 bind 创建上下文变量",
        "支持 function 定义、传参、与 foreach 组合及递归调用",
        "支持字符串、数字、数组、Map、布尔值和 nil 字面量",
        "支持属性访问与方法调用",
        "支持算术、数字分隔、比较、逻辑、字符串、成员、范围和三元运算符",
        "支持 len、all、none、any、one、filter、map 内置函数",
        "支持闭包和数组/字符串切片",
        "${} 与 #{} 中的表达式均可求值，#{} 以参数化形式输出",
    ],
    "S14": [
        "搜索和筛选已上架 API",
        "查看 API 详情并进入调用、授权、报警、监控和测试",
        "跨项目申请 API 调用权限",
        "选择应用、最大 QPS、有效期和申请原因",
        "申请行级/列级权限",
        "通过项目上架策略和审批控制 API 入市",
    ],
    "S15": [
        "搜索同业务线其他项目的逻辑表",
        "查看字段、预览探查和授权信息",
        "申请逻辑表权限",
        "选择授权项目、申请内容和申请原因",
        "通过授权实现同业务线跨项目复用",
    ],
    "S16": [
        "管理系统管理员、业务线管理员和项目管理员",
        "管理数据开发、API 开发、API 发布员和 QA 角色",
        "按租户、业务线和项目控制角色可见范围",
        "添加账户权限并选择账号、角色、业务线和项目",
        "由普通用户申请项目角色权限",
        "搜索、编辑和删除账户权限",
        "通过审批中心完成审核、状态查看和申请撤销",
    ],
    "S17": [
        "管理员直接添加项目，其他角色提交创建申请",
        "配置项目名称、描述、业务线和项目管理员",
        "设置 API 发布、上架、下线和调大 QPS 的审批规则",
        "控制 API 自动上架开关",
        "搜索、编辑和删除项目",
        "维护项目成员及其角色",
        "按项目批量授权应用调用 API",
        "配置项目封禁时间、白名单、状态和页面提示",
        "通过审批中心完成项目创建审核或撤销",
    ],
    "S18": [
        "创建业务线并指定业务线管理员和描述",
        "搜索、编辑和删除业务线",
        "在业务线下承载多个项目并支持经审批的资源复用",
        "兼容指标平台同步生成的业务线/项目并提示跨平台命名与删除风险",
    ],
    "S19": [
        "创建应用并维护全局唯一三段式标志符和管理员",
        "搜索、编辑和删除应用",
        "为一个应用创建多个密钥",
        "支持静态密钥和 OAuth 2.0 动态 Token",
        "配置密钥有效期、AppKey、AppSecret 和描述",
        "编辑、查看或删除密钥，并在 API 调用时校验密钥有效性",
    ],
    "S20": [
        "准备未绑定公网 IP、私有网络和辅助网卡",
        "将辅助网卡绑定公网 IP",
        "配置辅助网卡 ID、私网 ID 和公网 IP",
        "绑定后通过公网调用数据服务 API",
        "仅租户主账号可配置，部分地域不支持",
    ],
    "S21": [
        "配置一个或多个私网 ID",
        "按需生成 API 调用域名",
        "选择可用区子网和数据服务资源组",
        "完成 PrivateLink 跨服务授权与终端节点准备",
        "配置后在同一 VPC 内调用 API，不支持跨 VPC",
    ],
    "S22": [
        "我的审核按进行中、已完成、全部筛选工单",
        "查看审核详情并通过或驳回待审批工单",
        "我的申请按状态筛选并查看申请详情",
        "在审批前撤销自己的申请",
        "覆盖项目、权限、API 发布/上架/下线/QPS 等平台工单",
    ],
    "S23": [
        "创建公开或私有标签组",
        "维护标签组名称、描述和类型",
        "在标签组内新增、编辑和删除标签",
        "维护标签名称和描述",
        "关联可使用标签的项目",
        "搜索、编辑、复制和删除标签组",
        "在物理表、逻辑表和 API 中引用标签进行分类与检索",
    ],
}


PAGE_NOTES = {
    "S4": "网络方式决定白名单与安全组配置。EMR Doris/StarRocks 常用 8030、9030 端口；公网地址和地域信息应以当前官方页面为准。",
    "S5": "切流会直接影响线上服务。处于切流状态时修改目标数据源不生效，多数据源 Join 需重点检查跨集群风险。",
    "S6": "字段属性变更会直接生效并可能影响下游；历史版本回退属于应急能力。",
    "S7": "字段编辑、物理表切换和主备切换都可能影响下游 API；存在关联 API 时不能删除逻辑表。",
    "S8": "线上版本不可直接覆盖保存；线上发布、下线和调大 QPS 是否审批取决于项目策略。",
    "S9": "API 编排可按 API 使用，但当前官方页面明确编排工作流暂不支持灰度策略。",
    "S11": "调用监控当前页面说明的可查看范围为近 7 天；高敏业务不应只依赖默认失败率报警规则。",
    "S12": "OneService SQL 不以分号收尾；窗口函数和部分 MySQL 扩展语法不在当前支持范围。",
    "S13": "Dynamic SQL 的表达式不是底层 SQL 引擎表达式；XML 中的 <、>、&、引号等符号需要按场景转义。",
    "S16": "只有租户主账号、系统管理员、业务线管理员和项目管理员可按各自范围直接添加权限。",
    "S17": "指标平台同步生成的业务线/项目具有跨产品关联风险，不宜随意删除或单边改名。",
    "S19": "密钥类型创建后不可切换；OAuth 2.0 调用需先用 AppKey/AppSecret 获取动态 Token。",
    "S20": "仅租户主账号可配置；绑定后的官方服务网卡不能继续执行普通网卡操作。",
    "S21": "仅租户主账号可配置，且只支持同一 VPC 内调用，不支持跨 VPC。",
}


SECTION_OVERRIDES = {
    ("S1", "1 常见的业务问题"): "梳理服务化前的异构接入、权限可用性、运维和重复建设问题，并归纳效率、成本和质量痛点。",
    ("S1", "2 数据服务价值"): "用多种开发模式降低 API 生产门槛，并把版本、鉴权、转移、报警和监控纳入统一平台。",
    ("S1", "3 数据服务优势"): "统一接口与数据逻辑，减少重复计算/存储，并通过审核、鉴权和流控强化安全。",
    ("S3", "1 使用准备"): "先建立业务线、项目、人员角色和网络环境，作为后续数据资产与 API 的治理边界。",
    ("S3", "2 使用流程"): "按资产供给、模型映射、服务生产、发布授权和调用消费的顺序完成端到端交付。",
    ("S4", "3 添加数据源"): "选择数据源类型，录入连接信息，执行连通性测试后注册到平台。",
    ("S4", "4 数据源管理"): "提供搜索、编辑、切流和删除等数据源生命周期操作。",
    ("S5", "2 操作步骤"): "选择目标数据源、调用应用和受影响 API，设置切流比例并观察线上服务。",
    ("S6", "2 新建物理表"): "从数据源选择库表，确认字段治理属性和基本信息后注册物理表。",
    ("S6", "3 查看物理表"): "查看基础信息、字段明细和变更记录，并支持刷新、编辑与版本回退。",
    ("S6", "4 物理表管理"): "通过目录树搜索、定位、标签、负责人和删除操作维护物理表。",
    ("S7", "2 新建逻辑表"): "通过极简批量或标准单表模式建立物理表映射、字段模型、约束、授权和发布管控。",
    ("S7", "3 查看逻辑表"): "集中查看基础信息、字段、数据探查、授权和变更记录。",
    ("S7", "4 编辑逻辑表"): "维护字段与基本信息，切换来源物理表，设置主备，并按依赖关系控制删除。",
    ("S7", "5 目录树操作"): "支持搜索、目录、定位、移动、标签、负责人和删除。",
    ("S8", "2 文件夹管理"): "建立最多三级目录并维护重命名、移动和删除。",
    ("S8", "3 新建 API"): "选择开发模式并配置 API 标识、容量、安全、报警和版本元数据。",
    ("S8", "4 开发 API"): "按脚本式、向导式或原生式配置查询逻辑、参数与高级运行选项。",
    ("S8", "4.1 脚本式"): "编写普通或动态 SQL，绑定同源逻辑表，解析/维护请求与返回参数并支持分页方案。",
    ("S8", "4.2 向导式"): "通过界面选表与字段自动生成查询，无需手写 SQL。",
    ("S8", "4.3 原生式"): "由调用方在授权数据集范围内动态传入 SQL，适合分析面板类场景。",
    ("S8", "5 保存 API"): "支持更新当前版本、另存新版本或复制为新 API。",
    ("S8", "6 测试 API"): "在发布前通过 dryRun、表单/JSON 参数和结果/日志反馈验证逻辑。",
    ("S8", "7 发布API"): "发布到测试或线上环境；线上是否需要审批由项目策略决定。",
    ("S8", "8 基本配置"): "维护名称、负责人、QPS、等级、limit、描述和标签。",
    ("S8", "9 版本信息"): "查看历史版本、环境状态、差异、发布/下线/删除，并配置灰度。",
    ("S8", "10 API 详情"): "进入调用信息、授权管理、报警配置和调用监控。",
    ("S9", "3.3 API 编排开发"): "在画板上组合节点与依赖，定义入参、分支、调用、合并和出参。",
    ("S9", "3.3.1 界面操作说明"): "提供拖拽节点、连线依赖、缩放定位、自动布局和节点增删。",
    ("S9", "3.3.2 节点说明"): "定义开始、结束、函数、分支、API 与合并六类节点的职责。",
    ("S9", "3.4 编排测试"): "以开始节点参数执行整个编排，并把返回结果解析到结束节点。",
    ("S9", "3.5 编排发布"): "将工作流发布到测试或线上环境供外部调用。",
    ("S10", "2 操作流程"): "从 API 详情获取环境地址和代码示例，准备密钥、授权与网络后调用。",
    ("S10", "2.1 请求参数说明"): "覆盖应用密钥、授权、请求头、环境、API ID 和 Params/Sql 请求体。",
    ("S10", "2.2 返回参数说明"): "说明通用返回结构、字段类型和查询错误类型。",
    ("S11", "2 限流"): "在基本信息中开启最大 QPS 限流。",
    ("S11", "3 授权管理"): "向应用授权 API，并支持编辑有效期/QPS等权限或回收。",
    ("S11", "4 血缘关系"): "追踪逻辑表、API、应用和底层物理存储之间的依赖。",
    ("S11", "5 报警配置"): "维护报警规则、模板和告警接收通道。",
    ("S11", "6 调用监控"): "按应用、环境和时间查看 QPS、成功率和 PCT99。",
    ("S12", "1 占位符"): "区分直接插值与按类型参数化生成 SQL 片段的两类占位方式。",
    ("S12", "2 函数"): "提供日期格式、日期偏移和逻辑表就绪分区计算函数。",
    ("S13", "1 Dynamic SQL Engine 概述"): "说明以 XML 元素组织动态 SQL、表达式和参数求值的总体机制。",
    ("S13", "2 Dynamic SQL 元素"): "覆盖条件、选择、前后缀处理、遍历、变量绑定和函数元素。",
    ("S13", "3 表达式计算"): "覆盖字面量、对象访问、运算符、内置函数、闭包和切片。",
    ("S14", "4 操作步骤"): "在 API 集市搜索资产、查看详情并提交调用与行列权限申请。",
    ("S15", "3 操作步骤"): "搜索逻辑表、查看详情并提交跨项目授权申请。",
    ("S16", "1 角色权限说明"): "列出各层级管理员、开发、发布和 QA 角色的权限范围。",
    ("S16", "4 操作步骤"): "支持直接添加、申请、搜索、编辑和删除账户权限。",
    ("S17", "3 操作步骤"): "覆盖项目创建/申请、审批策略、成员、应用授权和封禁管理。",
    ("S18", "2 操作步骤"): "支持业务线新增、搜索、编辑和删除。",
    ("S19", "1 管理应用"): "维护应用名称、唯一标志符和管理员，并支持搜索、编辑与删除。",
    ("S19", "2 管理应用密钥"): "创建和维护静态或 OAuth 2.0 密钥，用于 API 调用鉴权。",
    ("S20", "3 操作步骤"): "绑定辅助网卡、私网和公网 IP。",
    ("S20", "4 操作结果"): "配置完成后允许通过公网调用已发布 API。",
    ("S21", "3 操作步骤"): "配置私网、域名、子网和资源组。",
    ("S21", "4 操作结果"): "配置完成后允许在同一 VPC 内调用已发布 API。",
    ("S22", "1 我的审核"): "查看待审/已审工单，进入详情并通过或驳回。",
    ("S22", "2 我的申请"): "查看申请进度和详情，并在审批前撤销。",
    ("S23", "2 操作步骤"): "维护标签组、标签及关联项目，并支持搜索、编辑、复制和删除。",
}


GENERIC_HEADINGS = {
    "前提条件": "列出使用该功能前必须满足的角色、项目、资产或网络条件。",
    "使用前提": "列出使用该功能前必须满足的角色、项目、资产或网络条件。",
    "使用前期": "列出调用或运维前必须完成的发布、应用、授权和网络准备。",
    "约束限制": "说明该功能的角色、地域、网络或对象边界。",
    "操作步骤": "按照控制台入口、配置、提交和结果确认的顺序执行。",
    "操作流程": "按照控制台入口、配置、提交和结果确认的顺序执行。",
    "适用场景": "说明该功能所解决的复用、共享或治理场景。",
    "后续操作": "说明提交后的审批、状态查看与撤销处理。",
    "语法": "列出函数或表达式的签名、输入参数和输出规则。",
    "示例": "覆盖字面量及占位符组合下的典型输入与生成结果。",
}


def clean_heading(text: str) -> str:
    return " ".join(text.replace("#", "").split())


def section_summary(sid: str, heading: str) -> str:
    heading = clean_heading(heading)
    if (sid, heading) in SECTION_OVERRIDES:
        return SECTION_OVERRIDES[(sid, heading)]
    base = heading
    while base and (base[0].isdigit() or base[0] in ". "):
        base = base[1:]
    for label, summary in GENERIC_HEADINGS.items():
        if base == label or base.endswith(label):
            return summary
    if "占位符" in base:
        return f"说明 {base} 的参数替换方式、类型处理和生成 SQL 行为。"
    if base in {"TO_YYYYMMDD", "ADD_DATE", "SUB_DATE", "DATE_PARTITION", "DATE_SUB_PARTITION"}:
        return f"定义 {base} 的函数签名、输入输出与占位符组合方式。"
    if base in {"if", "choose, when, otherwise", "trim, where", "foreach", "bind", "function"}:
        return f"说明 Dynamic SQL 元素 {base} 的用途、属性和生成 SQL 的行为。"
    if "Operators" in base or base in {"Supported Literals", "Builtin functions", "Closures", "Slices"}:
        return f"列出表达式引擎的 {base} 能力及典型用法。"
    if base.startswith("添加 "):
        return f"列出 {base[3:]} 的基础信息、网络/实例参数、账号和连接池配置。"
    return f"完整覆盖“{base}”对应的页面功能、操作和配置边界。"


def set_no_picture_compression(doc) -> None:
    settings = doc.settings.element
    node = settings.find(qn("w:doNotCompressPictures"))
    if node is None:
        node = OxmlElement("w:doNotCompressPictures")
        settings.append(node)
    node.set(qn("w:val"), "true")


def fit_image(width_px: int, height_px: int, max_width: float = 6.25, max_height: float = 6.05):
    aspect = width_px / height_px
    width = max_width
    height = width / aspect
    if height > max_height:
        height = max_height
        width = height * aspect
    return width, height


def add_original_image(doc, image_path: Path, alt: str, caption: str) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    width, height = fit_image(width_px, height_px)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
    set_picture_alt(paragraph, f"{alt}。原始分辨率 {width_px}×{height_px} 像素，未降采样。")
    add_caption(doc, f"{caption}（{width_px}×{height_px} px；官方原图直接嵌入，未降采样）")


def add_config_labels(doc, bullet_id: int, blocks: list[dict]) -> None:
    labels = []
    skip = {"参数", "配置项", "操作", "序号", "字段名称", "参数配置", "返回信息", "所属层级"}
    for block in blocks:
        if block.get("type") != "table":
            continue
        for row in block.get("rows", [])[1:]:
            if not row:
                continue
            label = " ".join(str(row[0]).split())
            if label and label not in skip and label not in labels:
                labels.append(label)
    if labels:
        shown = "、".join(labels)
        add_bullet(doc, bullet_id, shown, label="配置/能力项")


def add_code_coverage(doc, blocks: list[dict]) -> None:
    count = sum(1 for block in blocks if block.get("type") == "code")
    if count:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run("示例覆盖  ")
        set_font(run, size=9.6, bold=True, color=GREEN)
        run = paragraph.add_run(f"官方页面在本节提供 {count} 组语法/请求示例；本报告列全功能语义，代码原文请通过页面链接查看。")
        set_font(run, size=9.6, color=INK)


def page_sections(blocks: list[dict]) -> list[dict]:
    sections = []
    current = {"heading": None, "level": None, "blocks": []}
    for block in blocks:
        if block.get("type") == "heading":
            if current["heading"] is not None or current["blocks"]:
                sections.append(current)
            current = {"heading": block.get("text", ""), "level": block.get("level", 2), "blocks": []}
        else:
            current["blocks"].append(block)
    if current["heading"] is not None or current["blocks"]:
        sections.append(current)
    return sections


def add_page_source(doc, sid: str, url: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("官方来源：")
    set_font(run, size=9.2, bold=True, color=INK)
    add_hyperlink(paragraph, f"{sid} · {url}", url, size=9.2)


def set_extra_heading_styles(doc) -> None:
    specs = {
        "Heading 4": (10.8, DARK_BLUE, 7, 3),
        "Heading 5": (10.1, MUTED, 5, 2),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = hex_rgb(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_hierarchy_table(doc) -> None:
    rows = [
        ("数据服务简介", "概述；基本概念；快速入门", "—"),
        ("数据源", "创建数据源；数据源切流；物理表管理；逻辑表管理", "—"),
        ("API", "API 开发；API 编排开发；API 调用；API 运维；OneService 语法；Dynamic SQL 语法", "—"),
        ("数据集市", "API集市；逻辑表集市", "—"),
        ("系统管理", "账户权限管理；项目管理；业务线管理；应用管理；审批中心；标签管理", "网络配置 → 公网配置；VPC配置"),
    ]
    add_table(doc, ["一级导航", "二级页面", "三级页面"], rows, [1800, 4700, 2860])


def build() -> None:
    content = json.loads(CONTENT_JSON.read_text(encoding="utf-8"))
    inventory = json.loads(INVENTORY_JSON.read_text(encoding="utf-8"))
    page_lookup = {page["sid"]: page for page in content["pages"]}
    image_lookup = {}
    for page in inventory["pages"]:
        for image in page.get("images", []):
            if image.get("url") and image.get("localName"):
                image_lookup[image["url"]] = IMAGE_DIR / image["localName"]
                image_lookup[Path(image["url"].split("?", 1)[0]).name] = IMAGE_DIR / image["localName"]

    doc, bullet_id, _ = setup_document()
    set_extra_heading_styles(doc)
    set_no_picture_compression(doc)

    # Cover
    paragraph = doc.add_paragraph(style="Title")
    run = paragraph.add_run("火山引擎 DataLeap 数据服务")
    set_font(run, size=23, bold=True, color=INK)
    paragraph = doc.add_paragraph(style="Subtitle")
    run = paragraph.add_run("业务架构梳理 · 官网层级 · 全功能版")
    set_font(run, size=14, color=MUTED)
    add_callout(
        doc,
        "交付范围",
        "主文档严格按照官网“数据服务”目录展开 23 个页面，并继续沿用各页面正文的标题层级。功能、操作、配置项、语法能力与官方产品原图均按页面归档。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    add_body(doc, "资料来源：火山引擎《大数据研发治理套件》数据服务官方文档；页面核对日期：2026-08-25。")
    add_body(doc, "图片策略：保留用户提供的目录截图；正文官方产品图片全部按原始文件直接嵌入，仅调整显示尺寸，不降采样、不转码。")

    doc.add_page_break()
    doc.add_heading("阅读说明", level=1)
    add_bullet(doc, bullet_id, "官网一级导航作为 Heading 1；页面作为 Heading 2；网络配置下的页面作为 Heading 3。", label="目录规则")
    add_bullet(doc, bullet_id, "页面正文的 H2/H3/H4 顺延为文档 Heading 3/4/5，保留原始编号和名称。", label="正文规则")
    add_bullet(doc, bullet_id, "“完整功能清单”列出页面全部能力；后续按官方正文层级逐节说明，并附配置/能力项。", label="功能规则")
    add_bullet(doc, bullet_id, "官方截图证明页面/入口/操作可见，不等同于生产环境已经成功执行。", label="证据边界")
    add_bullet(doc, bullet_id, "语法页面列全函数、元素、运算符和内置能力，但不逐字复制全部示例代码。", label="版权与可用性")

    doc.add_heading("官网目录总览", level=1)
    add_hierarchy_table(doc)
    if SCREENSHOT.is_file():
        add_original_image(
            doc,
            SCREENSHOT,
            "用户提供的火山引擎文档页面截图，左侧红框标出数据服务目录层级",
            "图 1 用户提供的“数据服务”目录截图",
        )

    doc.add_heading("跨页面业务架构总览", level=1)
    if DIAGRAM.is_file():
        add_original_image(
            doc,
            DIAGRAM,
            "DataLeap 数据服务业务架构图，展示治理控制面、资产供给、逻辑建模、服务生产、集市复用和运行保障",
            "图 2 数据服务跨页面业务架构",
        )
    add_body(
        doc,
        "官网页面共同形成“治理边界 → 数据源接入 → 物理/逻辑资产 → API/编排生产 → 测试发布 → 授权调用 → 集市复用 → 运维监控”的主价值流。以下正文不再按抽象架构层拆分，而是严格回到官网页面层级。",
    )

    figure_no = 2
    page_counter = 0
    current_module = None
    network_heading_added = False
    meaningful_images = []

    for sid, name, module, page_id, _ in SOURCES:
        page = page_lookup[sid]
        top_module = module.split("/", 1)[0]
        if top_module != current_module:
            doc.add_page_break()
            doc.add_heading(top_module, level=1)
            current_module = top_module
            network_heading_added = False

        is_network = module == "系统管理/网络配置"
        if is_network and not network_heading_added:
            if page_counter:
                doc.add_page_break()
            doc.add_heading("网络配置", level=2)
            network_heading_added = True
        elif not is_network and page_counter:
            doc.add_page_break()

        page_level = 3 if is_network else 2
        doc.add_heading(f"{sid} {name}", level=page_level)
        add_page_source(doc, sid, page["url"])
        add_callout(doc, "页面定位", PAGE_OVERVIEWS[sid], fill=CALLOUT, accent=DARK_BLUE)
        if sid in PAGE_NOTES:
            add_callout(doc, "关键边界", PAGE_NOTES[sid], fill="FFF7E6", accent=GOLD)

        heading_level = min(5, page_level + 1)
        doc.add_heading("完整功能清单", level=heading_level)
        for item in PAGE_FUNCTIONS[sid]:
            add_bullet(doc, bullet_id, item)

        outline_rows = []
        for block in page["blocks"]:
            if block.get("type") == "heading":
                text = clean_heading(block.get("text", ""))
                outline_rows.append((f"H{block.get('level', 2)}", text, section_summary(sid, text)))
        if outline_rows:
            doc.add_heading("本页官方正文层级", level=heading_level)
            add_table(doc, ["官网层级", "官网标题", "功能覆盖"], outline_rows, [1050, 2870, 5440])

        sections = page_sections(page["blocks"])
        had_meaningful_image = False
        for section in sections:
            if section["heading"]:
                official_level = int(section["level"] or 2)
                level = min(5, page_level + official_level - 1)
                heading = clean_heading(section["heading"])
                doc.add_heading(heading, level=level)
                add_body(doc, section_summary(sid, heading))
            add_config_labels(doc, bullet_id, section["blocks"])
            add_code_coverage(doc, section["blocks"])

            for block in section["blocks"]:
                if block.get("type") != "image":
                    continue
                src = block.get("src", "")
                image_path = image_lookup.get(src) or image_lookup.get(Path(src.split("?", 1)[0]).name)
                if not image_path or not image_path.is_file():
                    raise FileNotFoundError(f"Missing original image for {sid}: {src}")
                with Image.open(image_path) as source_image:
                    width_px, height_px = source_image.size
                if width_px < 300 or height_px < 200:
                    continue
                figure_no += 1
                had_meaningful_image = True
                meaningful_images.append(image_path)
                context = clean_heading(section["heading"] or name)
                add_original_image(
                    doc,
                    image_path,
                    f"[{sid}] {name}页面“{context}”对应的火山引擎官方产品截图",
                    f"图 {figure_no} [{sid}] {name} · {context}",
                )

        if not had_meaningful_image:
            add_callout(
                doc,
                "截图状态",
                "当前官方正文未提供可作为产品界面的原始截图；本页仍按官网标题完整列出功能、配置和约束。",
                fill="FFF7E6",
                accent=GOLD,
            )
        page_counter += 1

    if page_counter != 23:
        raise AssertionError(f"Expected 23 pages, got {page_counter}")
    if len(meaningful_images) != 44:
        raise AssertionError(f"Expected 44 meaningful official images, got {len(meaningful_images)}")

    doc.core_properties.title = "火山引擎 DataLeap 数据服务业务架构梳理（官网层级·全功能版）"
    doc.core_properties.subject = "严格按照官方页面层级展开的数据服务业务架构与全功能说明"
    doc.core_properties.keywords = "火山引擎, DataLeap, 数据服务, 官网层级, 全功能, 官方截图"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Pages covered: {page_counter}")
    print(f"Official screenshots embedded: {len(meaningful_images)}")
    print(f"Total inline images: {len(meaningful_images) + 2}")


if __name__ == "__main__":
    build()
