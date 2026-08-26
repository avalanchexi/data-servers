from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TMP = ROOT / ".tmp" / "volcengine-data-service-report"
TMP.mkdir(parents=True, exist_ok=True)
OUTPUT = ROOT / "火山引擎DataLeap数据服务业务架构梳理.docx"
SCREENSHOT = Path(r"C:\Users\何峰\AppData\Local\Temp\codex-clipboard-9217ad3c-1a86-42ab-b592-2f20126fb448.png")
DIAGRAM = TMP / "data-service-business-architecture.png"

FONT_CN = r"C:\Windows\Fonts\msyh.ttc"
FONT_CN_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "C9D2DC"
WHITE = "FFFFFF"
GREEN = "2E7D5B"
GOLD = "A66A00"
RED = "9B1C1C"


SOURCES = [
    ("S1", "概述", "数据服务简介", 127499, "定位、价值、优势与‘配置即服务’"),
    ("S2", "基本概念", "数据服务简介", 127705, "业务线、项目、物理表、逻辑表、应用定义"),
    ("S3", "快速入门", "数据服务简介", 1154939, "从准备到调用的端到端主流程"),
    ("S4", "创建数据源", "数据源", 127696, "数据源注册、连接、连通性与支持类型"),
    ("S5", "数据源切流", "数据源", 127697, "数据库负载/异常下的流量迁移"),
    ("S6", "物理表管理", "数据源", 127698, "物理表元数据注册与安全等级"),
    ("S7", "逻辑表管理", "数据源", 1154905, "逻辑模型、查询约束、授权、变更与主备"),
    ("S8", "API 开发", "API", 127699, "开发模式、测试、发布、版本、灰度"),
    ("S9", "API 编排开发", "API", 1208340, "串并行编排及节点模型"),
    ("S10", "API 调用", "API", 127700, "HTTP 调用、应用密钥与请求返回约定"),
    ("S11", "API 运维", "API", 127583, "限流、授权、血缘、告警、监控"),
    ("S12", "OneService 语法", "API", 127701, "SQL 占位符、函数与语法边界"),
    ("S13", "Dynamic SQL 语法", "API", 1254733, "动态 SQL 模板与条件生成"),
    ("S14", "API集市", "数据集市", 1148775, "API 上架、搜索、申请调用与复用"),
    ("S15", "逻辑表集市", "数据集市", 1148776, "同业务线跨项目逻辑表复用"),
    ("S16", "账户权限管理", "系统管理", 1148766, "租户/业务线/项目分层角色权限"),
    ("S17", "项目管理", "系统管理", 1148767, "项目隔离边界及发布、上架、QPS 等策略"),
    ("S18", "业务线管理", "系统管理", 1148768, "业务域分组与跨项目复用边界"),
    ("S19", "应用管理", "系统管理", 127558, "消费应用、授权主体与密钥管理"),
    ("S20", "公网配置", "系统管理/网络配置", 127565, "公网入口与租户级配置"),
    ("S21", "VPC配置", "系统管理/网络配置", 1148772, "私网调用、终端节点与资源组"),
    ("S22", "审批中心", "系统管理", 1148770, "申请、审核、撤销与工单状态"),
    ("S23", "标签管理", "系统管理", 1148771, "资产分类、公开/私有标签组与项目关联"),
]


def hex_rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, size: float | None = None, bold: bool | None = None,
             color: str | None = None, italic: bool | None = None,
             name: str = "Calibri") -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = hex_rgb(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], indent: int = 120) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "222222",
                  size: float = 9.2, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], alignments: Sequence[WD_ALIGN_PARAGRAPH] | None = None) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    alignments = alignments or [WD_ALIGN_PARAGRAPH.LEFT] * len(headers)
    for idx, text in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], text, bold=True, color=INK, size=9.3, align=alignments[idx])
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            set_cell_text(row.cells[idx], str(text), align=alignments[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE,
                  underline: bool = True, size: float = 9.2):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz_cs)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(rfonts)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(txt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    r = paragraph.add_run("第 ")
    set_font(r, size=9, color=MUTED)
    for field_name in ("PAGE", "NUMPAGES"):
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {field_name} "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run = paragraph.add_run()
        run._r.extend((fld_char1, instr, fld_char2))
        if field_name == "PAGE":
            mid = paragraph.add_run(" / ")
            set_font(mid, size=9, color=MUTED)
    tail = paragraph.add_run(" 页")
    set_font(tail, size=9, color=MUTED)


def add_numbering(doc: Document, fmt: str, marker: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_bullet(doc: Document, bullet_id: int, text: str, label: str | None = None) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, bullet_id)
    if label:
        r = p.add_run(label)
        set_font(r, bold=True, color=INK)
    r = p.add_run(text)
    set_font(r)


def add_numbered(doc: Document, decimal_id: int, text: str, label: str | None = None) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, decimal_id)
    if label:
        r = p.add_run(label)
        set_font(r, bold=True, color=INK)
    r = p.add_run(text)
    set_font(r)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> object:
    p = doc.add_paragraph(style="Normal")
    if bold_lead:
        r = p.add_run(bold_lead)
        set_font(r, bold=True, color=INK)
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc: Document, label: str, text: str, fill: str = CALLOUT,
                accent: str = BLUE) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(label + "  ")
    set_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED)


def set_picture_alt(paragraph, description: str) -> None:
    nodes = paragraph._p.xpath(".//wp:docPr")
    if nodes:
        nodes[-1].set("descr", description)


def draw_multiline(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                   font: ImageFont.FreeTypeFont, fill: str, align: str = "center",
                   spacing: int = 8) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 34
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        lines.append(current)
    line_h = font.size + spacing
    total = line_h * len(lines) - spacing
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        width = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + (x2 - x1 - width) / 2 if align == "center" else x1 + 18
        draw.text((x, y), line, font=font, fill=(fill if fill.startswith("#") else "#" + fill))
        y += line_h


def rounded_box(draw, box, title, body, fill, outline, title_font, body_font):
    fill_color = fill if fill.startswith("#") else "#" + fill
    outline_color = outline if outline.startswith("#") else "#" + outline
    draw.rounded_rectangle(box, radius=24, fill=fill_color, outline=outline_color, width=4)
    x1, y1, x2, y2 = box
    title_h = 70
    draw.rounded_rectangle((x1, y1, x2, y1 + title_h), radius=24, fill=outline_color)
    draw.rectangle((x1, y1 + title_h - 24, x2, y1 + title_h), fill=outline_color)
    draw_multiline(draw, (x1, y1, x2, y1 + title_h), title, title_font, WHITE)
    draw_multiline(draw, (x1, y1 + title_h, x2, y2), body, body_font, INK)


def arrow(draw, start, end, color=BLUE, width=7):
    pil_color = color if color.startswith("#") else "#" + color
    draw.line((start, end), fill=pil_color, width=width)
    x1, y1 = start
    x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 20
    pts = [
        (x2, y2),
        (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6)),
        (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(pts, fill=pil_color)


def make_diagram() -> None:
    img = Image.new("RGB", (2400, 1500), "#" + WHITE)
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(FONT_CN_BOLD, 48)
    section_font = ImageFont.truetype(FONT_CN_BOLD, 31)
    body_font = ImageFont.truetype(FONT_CN, 27)
    tiny_font = ImageFont.truetype(FONT_CN, 24)
    draw.text((120, 60), "DataLeap 数据服务业务架构（基于官方文档归纳）", font=title_font, fill="#" + INK)

    rounded_box(draw, (90, 170, 440, 1230), "治理控制面",
                "租户\n↓\n业务线\n↓\n项目\n↓\n角色/权限\n\n审批中心\n标签管理\n网络配置",
                "F7F9FB", DARK_BLUE, section_font, body_font)

    rounded_box(draw, (560, 190, 920, 480), "数据供给",
                "MySQL / ByteHouse\nDoris / StarRocks 等\n\n注册数据源 + 连通性",
                "F7FAFF", BLUE, section_font, tiny_font)
    rounded_box(draw, (1010, 190, 1370, 480), "资产注册",
                "物理表元数据\n字段 / 主键 / 安全等级\n\n不直接供 API 使用",
                "F7FAFF", BLUE, section_font, tiny_font)
    rounded_box(draw, (1460, 190, 1820, 480), "逻辑建模",
                "逻辑表映射\n字段规范 / 查询约束\n授权 / 主备 / 变更",
                "F7FAFF", BLUE, section_font, tiny_font)
    arrow(draw, (920, 335), (1010, 335))
    arrow(draw, (1370, 335), (1460, 335))

    rounded_box(draw, (560, 590, 1030, 930), "服务生产",
                "脚本式 / 向导式 / 原生式\nDynamic SQL / OneService\n\n保存 → 测试 → 发布\n版本 / 对比 / 下线 / 灰度",
                "F7FBF9", GREEN, section_font, tiny_font)
    rounded_box(draw, (1120, 590, 1580, 930), "服务编排",
                "API、函数、条件、合并节点\n串行 / 并行工作流\n\n作为特殊 API 统一发布",
                "F7FBF9", GREEN, section_font, tiny_font)
    arrow(draw, (1640, 480), (800, 590), GREEN)
    arrow(draw, (1030, 760), (1120, 760), GREEN)

    rounded_box(draw, (1680, 590, 2240, 930), "共享与消费",
                "API 集市 / 逻辑表集市\n申请、审批、授权、复用\n\n应用 + 密钥（静态/OAuth2.0）\nHTTP 经公网或 VPC 调用",
                "FFFBF3", GOLD, section_font, tiny_font)
    arrow(draw, (1580, 760), (1680, 760), GOLD)

    rounded_box(draw, (560, 1050, 2240, 1320), "运行保障与闭环",
                "应用级授权与 QPS 限流  |  调用监控（QPS、成功率、PCT99）  |  报警与升级  |  血缘追溯\n数据源切流  |  逻辑表主备/物理表切换  |  API 版本与灰度  |  上线/下线通知和审批",
                "FFF8F8", RED, section_font, tiny_font)
    arrow(draw, (1960, 930), (1960, 1050), RED)
    arrow(draw, (800, 930), (800, 1050), RED)
    arrow(draw, (440, 700), (560, 700), DARK_BLUE)

    draw.text((90, 1400), "说明：左侧为跨域治理能力；右侧为从数据资产到服务消费的主价值流；底部为贯穿全生命周期的运行保障。",
              font=tiny_font, fill="#" + MUTED)
    img.save(DIAGRAM, quality=95)


def setup_document() -> tuple[Document, int, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_rgb("222222")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    styles = {
        "Title": (23, INK, 0, 4),
        "Subtitle": (14, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in styles.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = hex_rgb(color)
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = hex_rgb(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r = hp.add_run("火山引擎 DataLeap | 数据服务业务架构")
    set_font(r, bold=True, color=INK, size=8.8)
    r = hp.add_run("\t官方文档梳理")
    set_font(r, color=MUTED, size=8.8)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    r = p.add_run("资料截止：2026-08-25")
    set_font(r, size=9, color=MUTED)
    tab = p.add_run("\t")
    set_font(tab, size=9)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    add_page_number(p)

    bullet_id = add_numbering(doc, "bullet", "•")
    decimal_id = add_numbering(doc, "decimal", "%1.")
    return doc, bullet_id, decimal_id


def add_source_marker(paragraph, marker: str) -> None:
    r = paragraph.add_run(f" [{marker}]")
    set_font(r, size=9, color=BLUE, bold=True)


def build_document() -> None:
    make_diagram()
    doc, bullet_id, decimal_id = setup_document()

    p = doc.add_paragraph(style="Title")
    r = p.add_run("火山引擎 DataLeap 数据服务业务架构梳理")
    set_font(r, size=23, bold=True, color=INK)
    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("基于“大数据研发治理套件 > 用户指南 > 数据服务”官方页面")
    set_font(r, size=14, color=MUTED)

    meta = [
        ("调研入口", "https://docs.volcengine.com/docs/6260/127499?lang=zh"),
        ("覆盖范围", "数据服务简介、数据源、API、数据集市、系统管理，共 23 个具体页面"),
        ("资料日期", "截至 2026-08-25；各页面更新时间不一，详见附录"),
        ("输出口径", "官方明确行为与架构归纳分开陈述"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + "：")
        set_font(r, bold=True, color=INK)
        if value.startswith("http"):
            add_hyperlink(p, "官方文档入口（概述）", value, size=10.8)
        else:
            r = p.add_run(value)
            set_font(r, size=10.8)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    add_callout(
        doc,
        "一句话结论",
        "数据服务不是单一的 API 生成器，而是一套以项目为访问控制边界、以逻辑表为数据契约、以 API/编排为服务产品、以应用为消费身份，并由审批、网络、监控、告警、血缘和切流共同保障的“数据服务生产与运营体系”。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )

    doc.add_heading("核心结论", level=1)
    for label, text in [
        ("主价值流：", "数据源 → 物理表 → 逻辑表 → API/编排 → 发布授权 → 应用调用 → 运行监控。"),
        ("组织边界：", "租户负责全局治理，业务线承载同类业务资源复用，项目承担协作隔离与访问控制，应用承担运行时消费身份。"),
        ("核心解耦点：", "API 必须使用逻辑表，不直接绑定物理表；逻辑表屏蔽存储差异，并承载字段规范、安全等级、查询约束、授权和主备。"),
        ("运行闭环：", "上线审批、版本、灰度、QPS、授权、告警、调用监控、血缘、数据源切流共同覆盖服务上线后的变更与稳定性。"),
    ]:
        add_bullet(doc, bullet_id, text, label)

    p = add_body(doc, "阅读提示：文中“官方明确”表示页面直接描述；“架构归纳”表示根据多个页面关系整理出的业务架构判断。")
    p.paragraph_format.space_before = Pt(4)
    doc.add_page_break()

    doc.add_heading("1 文档范围与目录依据", level=1)
    add_body(doc, "本报告以用户提供的概述页为入口，按左侧“数据服务”目录读取 5 个一级业务域。截图中的网页操作说明仅作为资料内容，不视为对本报告生成过程的指令。")
    add_table(
        doc,
        ["一级业务域", "页面数", "回答的核心问题", "核心产物/对象"],
        [
            ("数据服务简介", "3", "为什么需要数据服务，完整流程是什么", "业务线、项目、物理表、逻辑表、应用"),
            ("数据源", "4", "如何把异构存储纳入平台并形成稳定数据契约", "数据源、物理表、逻辑表、切流关系"),
            ("API", "6", "如何生产、发布、调用和运营数据服务", "API、API 编排、版本、授权、监控"),
            ("数据集市", "2", "如何发现、申请和复用已有资产", "API 商品、逻辑表共享、申请工单"),
            ("系统管理", "8", "如何组织、授权、联网、审批和分类", "业务线、项目、角色、应用、网络、审批、标签"),
        ],
        [1600, 900, 3560, 3300],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(SCREENSHOT), width=Inches(6.35))
    set_picture_alt(p, "火山引擎大数据研发治理套件文档截图，左侧红框标出数据服务目录，包括数据服务简介、数据源、API、数据集市和系统管理。")
    add_caption(doc, "图 1  用户提供的官方文档截图：左侧红框为本次梳理的“数据服务”目录范围（原图保留）")
    add_callout(doc, "范围说明", "“最佳实践”和“常见问题”下另有数据服务页面，但截图红框和用户请求指向“用户指南 > 数据服务”主目录，本报告将它们视为延伸材料，不纳入 23 页主范围。", fill=CALLOUT)

    doc.add_heading("2 总体业务架构", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(DIAGRAM), width=Inches(6.35))
    set_picture_alt(p, "DataLeap 数据服务业务架构图：治理控制面贯穿数据供给、资产注册、逻辑建模、服务生产、服务编排、共享消费和运行保障。")
    add_caption(doc, "图 2  DataLeap 数据服务总体业务架构（架构归纳）")

    add_callout(doc, "架构归纳", "整体可分为五层：组织与治理、数据资产与逻辑模型、服务生产、共享与消费、运行保障。官方目录按功能模块组织，本报告按业务价值流重新组合。", fill=LIGHT_BLUE)

    doc.add_heading("2.1 五层架构", level=2)
    layer_rows = [
        ("1. 组织与治理层", "租户、业务线、项目、角色、审批、标签、网络", "建立所有权、隔离、复用和合规边界"),
        ("2. 数据资产与逻辑模型层", "数据源、物理表、逻辑表", "把异构存储转换为稳定、可治理的数据契约"),
        ("3. 服务生产层", "API、Dynamic SQL、OneService、API 编排、测试、发布", "将数据契约封装为可调用的服务产品"),
        ("4. 共享与消费层", "API 集市、逻辑表集市、应用、密钥、网络入口", "完成发现、申请、授权和实际调用"),
        ("5. 运行保障层", "限流、监控、告警、血缘、版本、灰度、切流", "让服务在变更、故障和增长中保持可控"),
    ]
    add_table(doc, ["层次", "核心能力", "业务目的"], layer_rows, [2100, 3560, 3700])

    doc.add_heading("2.2 控制面与运行面", level=2)
    add_body(doc, "架构上存在两条相互配合的链路：")
    add_bullet(doc, bullet_id, "定义谁能创建、发布、审批、共享和配置网络，并通过业务线、项目、角色、标签和审批中心进行治理。", "控制面：")
    add_bullet(doc, bullet_id, "将数据查询封装为 API，经应用授权和密钥鉴权后，通过公网或 VPC 被业务系统调用，并持续产生监控与告警数据。", "运行面：")
    add_body(doc, "这一区分是架构归纳。官方页面明确提供了对应对象和功能，但没有把它们统一命名为“控制面/运行面”。")

    doc.add_heading("3 端到端业务流程", level=1)
    add_callout(doc, "官方主流程", "快速入门给出的链路是：准备业务线/项目/权限/网络 → 配置数据源 → 创建物理表 → 创建逻辑表 → 创建、开发、测试、上线 API → 授权应用 → 调用 API。 [S3]", fill=CALLOUT)
    steps = [
        ("治理准备：", "租户或业务线管理员创建业务线与项目，分配数据开发/API 开发/API 发布等角色，并配置公网或 VPC。"),
        ("数据源注册：", "绑定目标数据库并完成网络和连通性测试，使平台能够读取元数据和执行查询。"),
        ("物理表注册：", "把存储中的表和字段元数据登记到平台，补充主键、安全等级等属性。"),
        ("逻辑建模：", "基于物理表创建逻辑表，进行字段规范、类型转换、查询约束、安全等级、授权和主备设置。"),
        ("服务开发：", "基于逻辑表选择向导式、脚本式或原生式开发 API；复杂流程可使用 API 编排。"),
        ("测试与发布：", "在本地测试或测试环境验证后发布；线上发布可按项目策略触发审批、通知，并形成版本记录。"),
        ("授权与调用：", "将已发布 API 授权给应用，配置有效期和最大 QPS；应用使用静态密钥或 OAuth2.0 Token 通过 HTTP 调用。"),
        ("运营闭环：", "根据 QPS、成功率、PCT99、失败率等监控，配置告警；通过版本、灰度、下线、切流和血缘分析处理变更与故障。"),
    ]
    for label, text in steps:
        add_numbered(doc, decimal_id, text, label)

    doc.add_heading("4 核心业务对象及关系", level=1)
    object_rows = [
        ("租户", "全局", "包含业务线、系统管理员、网络配置", "全局治理与最高权限边界"),
        ("业务线", "租户", "包含多个项目；同线资源可审批复用", "把同类业务聚合为共享边界"),
        ("项目", "业务线", "成员、数据源、表、API、应用授权的管理单元", "多用户隔离和访问控制的主要边界"),
        ("账户角色", "租户/业务线/项目", "赋予创建、测试、发布、授权、审批等能力", "实现职责分离与最小权限"),
        ("数据源", "项目/治理范围", "连接具体数据库或计算存储", "把异构存储接入统一服务出口"),
        ("物理表", "数据源", "注册存储中的真实表及字段元数据", "提供运行时查询所需的底层元数据"),
        ("逻辑表", "物理表映射", "API 的必选数据契约，可授权、变更、主备", "屏蔽存储细节并稳定下游接口"),
        ("API", "项目", "使用逻辑表；经历开发、测试、发布、版本和运维", "可交付、可授权的数据服务产品"),
        ("API 编排", "项目", "组合 API、函数、条件、合并等节点", "承载跨节点的复杂服务逻辑"),
        ("API 版本/灰度策略", "API", "版本对比、发布、下线、流量比例", "降低线上变更风险"),
        ("应用", "企业消费侧", "获得 API 授权并持有一个或多个密钥", "运行时调用身份和责任主体"),
        ("授权关系", "API↔应用/逻辑表↔项目", "有效期、最大 QPS、申请与回收", "控制谁可以消费什么以及消费强度"),
        ("审批工单", "治理控制面", "项目/角色/API 发布、上架、下线、QPS 等", "把高风险变更纳入可追踪流程"),
        ("标签", "物理表/逻辑表/API", "公开或私有标签组，并关联项目", "补充文件夹无法表达的分类维度"),
        ("集市条目", "共享层", "已上架 API 或可申请逻辑表", "减少重复建设并促进跨项目复用"),
    ]
    add_table(doc, ["对象", "归属/层级", "核心关系", "业务意义"], object_rows, [1500, 1800, 3160, 2900])
    add_callout(doc, "关键关系", "数据源 1:N 物理表；物理表经映射形成逻辑表；API 只使用逻辑表，不直接使用物理表；API 与应用通过授权形成 N:N 消费关系；业务线与项目定义资产共享和隔离范围。 [S2][S7][S10][S16-S19]", fill=LIGHT_BLUE)

    doc.add_heading("5 各业务域架构拆解", level=1)
    doc.add_heading("5.1 系统管理：先建立治理边界，再生产服务", level=2)
    add_body(doc, "系统管理不是外围配置，而是业务架构的控制面。它决定资源属于谁、谁可以操作、变更是否需审批、应用如何被识别，以及调用链路从哪里进入。")
    system_rows = [
        ("业务线管理", "组织同类业务；一个业务线含多个项目", "同业务线资源经审批可跨项目复用"),
        ("项目管理", "基本协作与隔离单元", "可配置 API 发布、上架、下线、QPS 等治理策略"),
        ("账户权限", "系统管理员、业务线管理员、项目管理员、数据开发、API 开发、API 发布、QA", "将建模、开发、测试、发布和授权职责分离"),
        ("应用管理", "应用标志符、管理员、密钥、API 权限", "把 API 调用绑定到可识别的消费主体"),
        ("网络配置", "公网或 VPC；VPC 域名可依赖终端节点和数据服务资源组", "控制调用面的网络可达性"),
        ("审批中心", "我的审核/我的申请及进行中、已完成、终止", "承载所有需人工决策的工单"),
        ("标签管理", "公开/私有标签组，最多关联多个项目和对象", "提供跨文件夹的资产分类维度"),
    ]
    add_table(doc, ["模块", "官方明确能力", "架构作用"], system_rows, [1800, 3900, 3660])

    doc.add_heading("5.2 数据源：从存储接入到稳定数据契约", level=2)
    add_bullet(doc, bullet_id, "数据源注册需要连接信息、网络白名单/安全组和连通性测试；当前页面覆盖 MySQL、ByteHouse、Doris、StarRocks 等类型。", "接入：")
    add_bullet(doc, bullet_id, "平台登记真实表元数据，并补充字段、主键、安全等级等属性。", "物理表：")
    add_bullet(doc, bullet_id, "逻辑表映射物理表，承担字段标准化、查询约束、安全等级、项目授权、变更记录和主备。API 必须基于逻辑表。", "逻辑表：")
    add_bullet(doc, bullet_id, "当源数据库过载或异常时，项目管理员可按应用/API 选择目标数据源和切流比例；切流直接影响线上服务。", "切流：")
    add_callout(doc, "架构判断", "逻辑表是整个体系最关键的解耦层：上游存储可以切换或形成主备，下游 API 继续面向稳定的逻辑字段与约束。", fill=LIGHT_BLUE)

    doc.add_heading("5.3 API：把查询逻辑产品化", level=2)
    api_rows = [
        ("向导式", "低代码配置查询条件和返回字段", "标准查询、较低开发门槛"),
        ("脚本式", "自定义 SQL；支持同源多逻辑表和 Dynamic SQL", "复杂条件、聚合和动态逻辑"),
        ("原生式", "面向底层引擎的原生 SQL 能力", "需要更直接控制执行语句的场景"),
        ("API 编排", "组合 API、函数、条件、合并等节点，支持串并行", "跨 API/函数的复杂加工与流程编排"),
    ]
    add_table(doc, ["生产方式", "主要能力", "适用场景"], api_rows, [1700, 4300, 3360])
    add_body(doc, "API 生命周期由保存、测试、发布、版本、灰度、下线和详情运维构成。测试环境可直接发布；线上环境是否审批取决于项目配置。发布后可按应用授权、设置 QPS、查看血缘、配置报警和调用监控。 [S8-S11]")
    add_callout(doc, "开发安全边界", "OneService/Dynamic SQL 支持参数占位和条件生成。官方 Dynamic SQL 页面建议优先使用 #{} 以按数据类型生成 SQL 片段并降低 SQL 注入风险；语法能力存在明确限制，应将它视为平台 SQL 运行时而非任意数据库客户端。 [S12][S13]", fill=CALLOUT)

    doc.add_heading("5.4 数据集市：从‘资产存在’到‘资产可复用’", level=2)
    add_bullet(doc, bullet_id, "已上架 API 可被搜索、查看详情和申请调用；申请时绑定应用、最大 QPS、有效期等。", "API 集市：")
    add_bullet(doc, bullet_id, "同一业务线下，其他项目的逻辑表可被搜索、探查和申请授权；仅能为自己已加入的项目申请。", "逻辑表集市：")
    add_body(doc, "两类集市分别服务两种复用：API 集市复用“已封装的服务产品”，逻辑表集市复用“可继续开发的标准数据契约”。这是减少 API 和逻辑表重复开发的主要机制。")

    doc.add_heading("5.5 运行保障：把服务上线后的问题纳入平台", level=2)
    ops_rows = [
        ("容量与流控", "API 最大 QPS、应用授权 QPS、数据源限流", "防止单一消费者或 API 压垮存储"),
        ("可观测性", "QPS、成功率、请求失败率、PCT99、失败次数", "按应用和时间范围定位性能与稳定性问题"),
        ("告警", "Notice/P2/P1/P0、升级策略、短信/Webhook", "把异常推送给项目成员并形成响应升级"),
        ("变更控制", "版本记录、版本对比、发布/下线审批、通知、灰度比例", "降低发布和回滚风险"),
        ("依赖追溯", "数据源→物理表→逻辑表→API→应用血缘", "支持影响分析、下线决策和问题定位"),
        ("容灾", "逻辑表主备、物理表切换、数据源按比例切流", "在底层存储异常或迁移时维持服务"),
    ]
    add_table(doc, ["保障主题", "官方能力", "业务效果"], ops_rows, [1700, 4200, 3460])

    doc.add_heading("6 角色与职责架构", level=1)
    role_rows = [
        ("系统管理员", "租户", "管理业务线及租户内资源", "全局治理"),
        ("业务线管理员", "业务线", "管理业务线、项目及线内资源", "业务域治理和共享"),
        ("项目管理员", "项目", "注册数据源、管理成员和项目设置", "项目边界与策略负责人"),
        ("数据开发", "项目", "创建/编辑物理表、逻辑表和自己负责的 API", "数据契约生产"),
        ("API 开发", "项目", "读取有权限逻辑表，创建、测试和编辑自己的 API", "服务逻辑生产"),
        ("API 发布", "项目", "测试、发布、下线及应用授权", "发布与运行权限"),
        ("QA", "项目", "API 测试及特定环境表配置", "质量验证"),
        ("应用管理员", "应用", "管理应用及密钥、申请/查看 API 权限", "消费侧身份与凭据"),
    ]
    add_table(doc, ["角色", "作用域", "主要职责", "架构定位"], role_rows, [1600, 1400, 3860, 2500])
    add_callout(doc, "职责分离", "页面将 API 开发、测试、发布、下线和授权拆分给不同角色，并可叠加项目级审批策略。这说明平台的目标不仅是提高开发效率，也包括控制生产变更权限。 [S16][S17]", fill=LIGHT_BLUE)

    doc.add_heading("7 典型业务场景", level=1)
    scenarios = [
        ("场景 A：快速把数据库表服务化", "准备项目与权限 → 注册 MySQL/ByteHouse 等数据源 → 登记物理表 → 创建逻辑表 → 选择向导式或脚本式 API → 测试、发布 → 授权应用 → HTTP 调用。"),
        ("场景 B：同业务线跨项目复用", "先在逻辑表集市申请数据契约，或在 API 集市直接申请已发布服务；审批通过后绑定目标项目或应用，避免重复建表和重复开发接口。"),
        ("场景 C：复杂服务组合", "把多个 API、FaaS 函数、条件和合并节点编排成串行/并行流程，整体作为特殊 API 测试、发布和运维。"),
        ("场景 D：线上数据库异常", "利用调用监控和告警发现问题，通过血缘识别受影响 API/应用，再按应用和 API 配置数据源切流；必要时结合逻辑表主备和 API 灰度控制影响面。"),
        ("场景 E：高风险版本变更", "在测试环境验证，线上发布按项目策略审批并通知下游；先以灰度比例导入小部分流量，观察成功率/PCT99/告警，再扩大或回退。"),
    ]
    for title, body in scenarios:
        doc.add_heading(title, level=2)
        add_body(doc, body)

    doc.add_heading("8 架构优势、边界与关注点", level=1)
    doc.add_heading("8.1 官方明确的优势", level=2)
    for text in [
        "统一接口标准和服务元信息，屏蔽不同中间存储，减少下游对接。",
        "数据计算逻辑收敛到平台，促进跨应用口径一致和逻辑复用。",
        "减少应用侧重复计算、重复存储及敏感明细数据暴露。",
        "通过审核、鉴权、限流、监控和告警降低服务治理成本。",
    ]:
        add_bullet(doc, bullet_id, text)

    doc.add_heading("8.2 从页面关系推导的架构关注点", level=2)
    concern_rows = [
        ("逻辑表治理是核心", "字段、查询约束、授权或物理表映射变更都会影响下游 API", "把逻辑表视为正式数据契约，建立变更评审和影响分析"),
        ("项目是强边界", "大多数资源、角色和策略以项目为单位", "项目划分应贴合团队责任、业务域和数据访问边界"),
        ("应用是运行时责任主体", "授权、密钥、QPS 和监控均可落到应用", "统一应用标志符命名，定期轮换密钥和回收闲置权限"),
        ("治理策略可配置", "发布、上架、下线、QPS 审批取决于项目设置", "按业务等级设定差异化策略，避免一刀切或完全放开"),
        ("容灾能力分层", "数据源切流、逻辑表主备、API 灰度分别作用于不同层", "明确触发条件、负责人和回退路径，避免临时操作放大风险"),
        ("文档能力会演进", "页面更新时间跨度较大，部分入口和数据源类型已更新", "落地前应按当前租户控制台和最新页面再次核验"),
    ]
    add_table(doc, ["关注点", "依据", "建议"], concern_rows, [1800, 3400, 4160])

    doc.add_heading("9 业务架构总结", level=1)
    add_callout(doc, "最终判断", "DataLeap 数据服务的核心不是把 SQL 暴露成 HTTP，而是把数据供给、逻辑契约、服务产品、消费身份和运行治理放进同一生命周期。其业务架构以项目为隔离单元、以业务线为复用边界、以逻辑表为解耦层、以应用授权为消费控制点，并通过审批与可观测性形成运营闭环。", fill=LIGHT_BLUE, accent=DARK_BLUE)
    add_body(doc, "如果将这套能力映射到企业数据中台，可将其理解为“数据服务工厂 + 服务目录 + API 运营平台 + 数据访问治理”的组合：前半段负责生产标准服务，后半段负责让服务被安全、稳定、可追踪地消费。")

    doc.add_heading("附录 A 目录页清单与来源", level=1)
    add_body(doc, "以下页面均属于本次主范围。链接指向火山引擎官方文档，访问日期为 2026-08-25。")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [650, 2300, 1850, 4560])
    for idx, title in enumerate(["编号", "官方页面", "目录模块", "用于本报告的证据"]):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], title, bold=True, color=INK, size=9.3)
    set_repeat_table_header(table.rows[0])
    for sid, title, module, doc_id, purpose in SOURCES:
        row = table.add_row()
        set_cell_text(row.cells[0], sid, align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[1].text = ""
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_hyperlink(p, title, f"https://docs.volcengine.com/docs/6260/{doc_id}?lang=zh", size=9.2)
        set_cell_text(row.cells[2], module)
        set_cell_text(row.cells[3], purpose)

    doc.add_heading("附录 B 官方明确与架构归纳的边界", level=1)
    boundary_rows = [
        ("官方明确", "目录、对象定义、角色权限、配置项、使用前提、操作流程、监控指标、审批和切流规则", "可在对应来源页直接找到"),
        ("架构归纳", "五层架构、控制面/运行面、对象关系图、典型场景串联、落地关注点", "由多个官方页面之间的关系推导"),
        ("未覆盖", "具体租户的实际配置、组织职责、流量规模、SLA、成本与部署拓扑", "需要结合企业现状另行调研"),
    ]
    add_table(doc, ["类型", "内容", "证据口径"], boundary_rows, [1500, 5000, 2860])

    doc.core_properties.title = "火山引擎 DataLeap 数据服务业务架构梳理"
    doc.core_properties.subject = "基于火山引擎官方文档的数据服务业务架构分析"
    doc.core_properties.keywords = "火山引擎, DataLeap, 数据服务, 业务架构, API, 逻辑表"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUTPUT)


if __name__ == "__main__":
    if not SCREENSHOT.exists():
        raise FileNotFoundError(SCREENSHOT)
    build_document()
    print(OUTPUT)
