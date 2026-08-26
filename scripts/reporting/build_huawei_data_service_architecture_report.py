from __future__ import annotations

import json
import math
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_volcengine_data_service_report import (
    BLUE,
    CALLOUT,
    DARK_BLUE,
    GOLD,
    GREEN,
    INK,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    RED,
    add_body,
    add_bullet,
    add_callout,
    add_caption,
    add_hyperlink,
    add_numbered,
    add_numbering,
    add_page_number,
    add_source_marker,
    add_table as _base_add_table,
    hex_rgb,
    set_font,
    set_picture_alt,
)


ROOT = Path(__file__).resolve().parents[2]
TMP = ROOT / ".tmp" / "huawei-data-service-report"
INVENTORY_PATH = TMP / "inventory.json"
IMAGE_DIR = TMP / "page-images"
DIAGRAM = TMP / "dataartsstudio-data-service-application-architecture.png"
OUTPUT = ROOT / "华为云DataArts Studio数据服务应用架构梳理（含官网截图）.docx"
SOURCE_ROOT = "https://support.huaweicloud.com/intl/zh-cn/dataartsstudio/index.html"
RESEARCH_DATE = "2026-08-25"


SOURCES = [
    ("S01", "数据服务简介", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0301.html"),
    ("S02", "规格说明", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0315.html"),
    ("S03", "管理专享版集群", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0314.html"),
    ("S04", "配置方式生成API", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0305.html"),
    ("S05", "脚本/MyBatis方式生成API", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0306.html"),
    ("S06", "调试API", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0316.html"),
    ("S07", "发布API", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0308.html"),
    ("S08", "API版本管理", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0339.html"),
    ("S09", "API编排简介", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0324.html"),
    ("S10", "调用API方式简介", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0309.html"),
    ("S11", "通过应用授权APP认证方式API", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0333.html"),
    ("S12", "申请API授权", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0334.html"),
    ("S13", "配置API调用流控策略", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0311.html"),
    ("S14", "数据服务支持的监控指标", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_7690.html"),
    ("S15", "查看数据服务监控指标", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_7692.html"),
    ("S16", "查看API访问日志", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0900.html"),
    ("S17", "配置数据服务审核中心", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0312.html"),
    ("S18", "数据服务API概览", "https://support.huaweicloud.com/intl/zh-cn/api-dataartsstudio/apiOverview_DataService.html"),
    ("S19", "创建API - CreateApi", "https://support.huaweicloud.com/intl/zh-cn/api-dataartsstudio/CreateApi.html"),
    ("S20", "设置API可见", "https://support.huaweicloud.com/intl/zh-cn/usermanual-dataartsstudio/dataartsstudio_01_0323.html"),
]


def add_table(doc, headers, rows, widths, alignments=None):
    """Create a table whose rows remain intact across Word page breaks."""
    table = _base_add_table(doc, headers, rows, widths, alignments)
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if index == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    return table


def setup_document() -> tuple[Document, int, int]:
    """Resolve the standard_business_brief preset with an editorial-cover override."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_rgb("222222")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = {
        "Title": (27, INK, 0, 6, True),
        "Subtitle": (14, MUTED, 0, 16, False),
        "Heading 1": (16, BLUE, 16, 8, True),
        "Heading 2": (13, BLUE, 12, 6, True),
        "Heading 3": (12, DARK_BLUE, 8, 4, True),
    }
    for name, (size, color, before, after, bold) in style_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = hex_rgb(color)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = hex_rgb(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    run = paragraph.add_run("华为云 DataArts Studio | 数据服务应用架构")
    set_font(run, bold=True, color=INK, size=8.8)
    run = paragraph.add_run("\t官网证据梳理")
    set_font(run, color=MUTED, size=8.8)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(f"资料截止：{RESEARCH_DATE}")
    set_font(run, size=9, color=MUTED)
    paragraph.add_run("\t")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    add_page_number(paragraph)

    bullet_id = add_numbering(doc, "bullet", "•")
    decimal_id = add_numbering(doc, "decimal", "%1.")
    return doc, bullet_id, decimal_id


def draw_text_wrapped(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
    line_gap: int = 10,
    center: bool = True,
) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 34
    lines: list[str] = []
    for paragraph in text.split("\n"):
        current = ""
        for char in paragraph:
            test = current + char
            if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = char
        lines.append(current)
    line_height = font.size + line_gap
    total_height = line_height * len(lines) - line_gap
    y = y1 + max(0, (y2 - y1 - total_height) / 2)
    for line in lines:
        width = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + (x2 - x1 - width) / 2 if center else x1 + 18
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def rounded_panel(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: str,
    accent: str,
    title_font: ImageFont.FreeTypeFont,
    body_font: ImageFont.FreeTypeFont,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=accent, width=4)
    header_height = 66
    draw.rounded_rectangle((x1, y1, x2, y1 + header_height), radius=22, fill=accent)
    draw.rectangle((x1, y1 + header_height - 22, x2, y1 + header_height), fill=accent)
    draw_text_wrapped(draw, (x1, y1, x2, y1 + header_height), title, title_font, "#FFFFFF", 6)
    draw_text_wrapped(draw, (x1, y1 + header_height, x2, y2), body, body_font, "#17324D", 7)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str) -> None:
    draw.line((start, end), fill=color, width=7)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 20
    points = [
        (x2, y2),
        (x2 - size * math.cos(angle - math.pi / 6), y2 - size * math.sin(angle - math.pi / 6)),
        (x2 - size * math.cos(angle + math.pi / 6), y2 - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(points, fill=color)


def make_architecture_diagram() -> None:
    image = Image.new("RGB", (2400, 1650), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    font_dir = Path(r"C:\Windows\Fonts")
    title_font = ImageFont.truetype(str(font_dir / "msyhbd.ttc"), 50)
    panel_title = ImageFont.truetype(str(font_dir / "msyhbd.ttc"), 30)
    body_font = ImageFont.truetype(str(font_dir / "msyh.ttc"), 25)
    small_font = ImageFont.truetype(str(font_dir / "msyh.ttc"), 23)
    draw.text((105, 55), "DataArts Studio 数据服务应用架构（基于官网功能重构）", font=title_font, fill="#0B2545")

    rounded_panel(
        draw,
        (70, 180, 430, 1370),
        "治理与控制面",
        "DataArts Studio实例\n工作空间 / 项目 / IAM\n\nAPI目录与可见范围\n审核人 / 审核单\n应用 / 授权 / 白名单\n版本 / 导入导出 / 同步\n流控策略 / 配额\n\nCloud Eye 监控\nLTS / OBS 日志\n告警与运营视图",
        "#F7F9FC",
        "#315A7D",
        panel_title,
        body_font,
    )

    layer_boxes = [
        ((540, 170, 2260, 360), "调用与消费层", "企业内部应用  |  合作伙伴/开发者  |  SDK  |  API工具  |  浏览器（无认证限定场景）", "#F7FBFF", "#2773B8"),
        ((540, 405, 2260, 600), "接入与安全层", "VPC终端节点 / 内网域名  |  EIP / 公网域名  |  HTTP/HTTPS  |  安全组与路由  |  APP签名 / IAM Token / 白名单", "#FFF9F1", "#B77717"),
        ((540, 645, 2260, 865), "服务运行层", "专享版集群  |  终端节点服务  |  ELB  |  多节点  |  API路由与调用  |  用户/应用/时间维度流控", "#F5FBF8", "#2B8A66"),
        ((540, 910, 2260, 1125), "服务生产层", "API目录  |  配置式API  |  脚本/MyBatis API  |  API参数和返回字段  |  在线调试  |  审核发布  |  可视化API编排", "#F8F6FF", "#6E55A3"),
        ((540, 1170, 2260, 1375), "数据访问层", "管理中心数据连接  |  表/字段元数据  |  SQL取数逻辑  |  DLI / DWS / MySQL / RDS / Hive / HBase / ClickHouse / Hetu / GBase / Doris / Oracle", "#FFF7F8", "#A84C5E"),
    ]
    for box, title, body, fill, accent in layer_boxes:
        rounded_panel(draw, box, title, body, fill, accent, panel_title, small_font)

    for upper, lower, color in [
        ((1400, 360), (1400, 405), "#2773B8"),
        ((1400, 600), (1400, 645), "#2B8A66"),
        ((1400, 865), (1400, 910), "#6E55A3"),
        ((1400, 1125), (1400, 1170), "#A84C5E"),
    ]:
        arrow(draw, upper, lower, color)
    arrow(draw, (430, 760), (540, 760), "#315A7D")

    draw.rounded_rectangle((540, 1430, 2260, 1570), radius=18, fill="#EEF3F8", outline="#72869A", width=3)
    draw_text_wrapped(
        draw,
        (560, 1436, 2240, 1564),
        "主数据流：数据源 → API取数逻辑 → 专享集群执行 → 认证/流控 → 用户应用\n主控制流：工作空间与目录 → 开发/调试 → 审核/发布 → 授权 → 监控/日志/变更",
        small_font,
        "#17324D",
        7,
    )
    image.save(DIAGRAM, quality=95)


def get_inventory() -> dict:
    return json.loads(INVENTORY_PATH.read_text(encoding="utf-8"))


def get_image(inventory: dict, page_id: str, one_based_index: int) -> tuple[Path, dict, dict]:
    page = next(page for page in inventory["pages"] if page["id"] == page_id)
    image_info = page["images"][one_based_index - 1]
    return IMAGE_DIR / image_info["file"], page, image_info


def add_figure(
    doc: Document,
    inventory: dict,
    page_id: str,
    image_index: int,
    caption: str | None = None,
    width: float = 6.25,
) -> None:
    image_path, page, image_info = get_image(inventory, page_id, image_index)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    source_caption = image_info.get("caption") or f"{page['title']}官方页面截图"
    set_picture_alt(paragraph, f"华为云 DataArts Studio 数据服务产品截图：{source_caption}。")
    add_caption(doc, caption or f"官网截图：{source_caption}（来源：{page['title']}）")


def add_source_body(doc: Document, text: str, marker: str) -> object:
    paragraph = add_body(doc, text)
    add_source_marker(paragraph, marker)
    return paragraph


def add_source_list(doc: Document, sources: Iterable[tuple[str, str, str]]) -> None:
    for marker, title, url in sources:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"[{marker}] {title}：")
        set_font(run, size=9.6, bold=True, color=INK)
        add_hyperlink(paragraph, url, url, size=9.0)


def style_layer_table(table) -> None:
    fills = ["EAF4FD", "FFF3DF", "E8F5EF", "F0EBFA", "FCECEF", "EEF3F8"]
    for index, row in enumerate(table.rows[1:]):
        fill = fills[index % len(fills)]
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), fill)


def build_document() -> None:
    inventory = get_inventory()
    make_architecture_diagram()
    doc, bullet_id, decimal_id = setup_document()

    # Editorial cover override for a long-form technical report.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(80)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("产品与技术架构报告")
    set_font(run, size=11, bold=True, color=GOLD)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("华为云 DataArts Studio\n数据服务应用架构梳理")
    set_font(run, size=27, bold=True, color=INK)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于国际站官方用户指南与 API 参考，含产品原图")
    set_font(run, size=14, color=MUTED)

    meta_rows = [
        ("调研入口", SOURCE_ROOT),
        ("官网范围", "用户指南 > 数据服务；共5层文档深度、52个页面"),
        ("原图资产", f"{inventory['unique_image_count']} 张官网原图；正文精选关键架构与操作界面"),
        ("资料日期", f"截至 {RESEARCH_DATE}，各页更新时间详见附录"),
        ("口径", "“官网明确”与“架构归纳”分开陈述，不把推断写成产品承诺"),
    ]
    table = add_table(doc, ["项目", "内容"], meta_rows, [1600, 7760])
    for row in table.rows[1:]:
        row.cells[0].paragraphs[0].runs[0].bold = True

    add_callout(
        doc,
        "一句话结论",
        "DataArts Studio 数据服务是一套以专享集群为运行时、以API为数据产品、以应用与授权为消费边界，并由审批、流控、监控、日志和版本管理保障的数据开放平台。",
        fill=LIGHT_BLUE,
        accent=DARK_BLUE,
    )
    doc.add_page_break()

    doc.add_heading("执行摘要", level=1)
    for label, text in [
        ("业务定位：", "把表数据和查询逻辑封装为统一 RESTful API，减少上层应用与数据库的直接耦合。"),
        ("核心分层：", "数据连接与取数逻辑 → API开发/编排 → 审核/发布 → 专享集群运行 → 应用授权与调用。"),
        ("两条主链：", "提供方完成开发、调试、发布和授权；调用方完成发现、申请授权和认证调用。"),
        ("运行边界：", "官网主流程以专享版为主，需先购买专享集群，再在工作空间内分配 API 配额。"),
        ("安全边界：", "入网由 VPC/EIP/域名/安全组控制，调用由 APP 签名、IAM Token 或无认证模式控制，再叠加审批、可见性和流控。"),
        ("适用场景：", "小批量、快速响应、稳定数据契约的业务查询与系统集成；不适合将大数据集作为批量文件式导出。"),
    ]:
        add_bullet(doc, bullet_id, text, label)
    add_source_body(doc, "以上结论主要来自数据服务简介、专享集群、调用方式与规格说明。", "S01/S02/S03/S10")

    doc.add_heading("阅读地图", level=2)
    add_table(
        doc,
        ["章节", "回答的问题", "主要读者"],
        [
            ("1 总体应用架构", "整体如何分层，数据流与控制流如何衔接", "架构师、产品负责人"),
            ("2 部署与网络", "专享集群部署在哪里，如何被内外网访问", "云架构师、运维"),
            ("3 核心对象模型", "API、应用、授权、集群与工作空间如何关联", "产品、研发、治理"),
            ("4 服务生产链路", "如何创建、调试、发布、版本化与编排 API", "API开发者"),
            ("5 服务消费与安全", "如何发现、授权、认证和调用", "应用开发者、安全"),
            ("6 治理与可观测", "如何审核、限流、监控、查日志", "治理、审核、运维"),
            ("7 开放接口与集成", "如何对控制面进行自动化与集成", "平台研发、集成工程师"),
            ("8 非功能边界与落地蓝图", "怎样选型，哪些约束要前置验证", "决策者、实施团队"),
        ],
        [1600, 4860, 2900],
    )

    doc.add_heading("1 总体应用架构", level=1)
    add_source_body(doc, "DataArts Studio 将数据服务定位为企业统一数据服务总线：把数据表快速生成数据 API，并管理 API 的发布、运维与对内对外开放。", "S01")
    add_source_body(doc, "官网同时明确了产品边界：这是将数据逻辑封装成 RESTful API 的小批量快速交互方案，大量数据开放应选用数据共享交换或其他方案。", "S01/S02")
    add_figure(doc, inventory, "dataartsstudio_01_0301", 1, "图 1  官方数据服务架构图：数据源与已有API经数据服务集群向用户/应用开放")

    doc.add_heading("1.1 分层架构（基于官网功能重构）", level=2)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(DIAGRAM), width=Inches(6.32))
    set_picture_alt(paragraph, "DataArts Studio 数据服务应用架构图，展示消费、安全、运行、服务生产、数据访问及治理控制面。")
    add_caption(doc, "图 2  DataArts Studio 数据服务分层应用架构（架构归纳，非官方原图）")
    add_callout(doc, "口径说明", "官方文档按功能菜单组织，并未直接命名以上六层。这一分层是将 52 个页面按数据流、控制流和运行边界重组后得出的架构视图。", fill=CALLOUT)

    layers = [
        ("1. 调用与消费层", "内部应用、合作伙伴、开发者、SDK、API 工具和受限浏览器调用", "把数据能力嵌入上层业务"),
        ("2. 接入与安全层", "VPC终端节点、EIP、内/公网域名、HTTPS、安全组、APP/IAM/无认证", "建立网络与身份边界"),
        ("3. 服务运行层", "专享版集群、终端节点服务、ELB、节点、路由、流控", "执行API请求并保护后端"),
        ("4. 服务生产层", "API目录、配置式、脚本/MyBatis式、调试、审核发布、API编排", "把数据逻辑产品化为服务契约"),
        ("5. 数据访问层", "管理中心数据连接、表字段、SQL取数、多种数据源", "让API在后端执行可治理的查询"),
        ("6. 治理与控制面", "工作空间、IAM、目录、应用、授权、审批、版本、可见性、监控与日志", "贯穿服务生命周期"),
    ]
    table = add_table(doc, ["架构层", "核心能力", "主要职责"], layers, [2200, 4500, 2660])
    style_layer_table(table)

    doc.add_heading("1.2 两个视角：API提供方与调用方", level=2)
    add_table(
        doc,
        ["角色", "主流程", "关键对象", "成功标准"],
        [
            ("API提供方", "集群准备 → 审核人 → 创建 → 调试 → 发布 → 管理/编排 → 流控 → 主动授权", "API目录、API、版本、审核单、流控策略", "API已发布、可见、可被授权和稳定调用"),
            ("API调用方", "服务目录发现 → 申请授权 → 审批 → 获取认证信息 → SDK/工具调用", "服务目录、应用、授权、AppKey/AppSecret或Token", "在授权期、认证与流控范围内得到预期数据"),
        ],
        [1450, 3800, 2360, 1750],
    )
    add_source_body(doc, "官网将数据服务的使用者明确分成 API 开放方和 API 调用方，两条链路在发布后的服务目录、应用授权和审核中心会合。", "S01/S11/S12/S17")

    doc.add_heading("1.3 平台总览是运营驾驶舱", level=2)
    add_body(doc, "总览页从 API 和 APP 两个视角统计已发布/开发中、授权关系、调用总量、成功/失败、合法/非法请求、调用趋势和 TOP5。这说明产品并不只关注开发，还把服务产品和消费应用作为运营对象。")
    add_figure(doc, inventory, "dataartsstudio_01_0301", 2, "图 3  API视角数据统计：发布、调用趋势、调用率/时长/次数 TOP5")
    add_figure(doc, inventory, "dataartsstudio_01_0301", 3, "图 4  APP视角数据统计：认证调用与应用消费表现")

    doc.add_heading("2 部署与网络架构", level=1)
    add_source_body(doc, "专享版集群位于资源租户区，使用 ELB 对集群节点做负载均衡。用户 VPC 内的终端节点通过终端节点服务访问集群；如开启公网入口，还可通过绑定在 ELB 上的 EIP 访问。", "S03")
    add_figure(doc, inventory, "dataartsstudio_01_0314", 1, "图 5  官方专享版集群网络架构：VPC终端节点、终端节点服务、ELB、节点与EIP")

    doc.add_heading("2.1 四种访问入口", level=2)
    add_table(
        doc,
        ["入口", "是否默认", "适用场景", "关键前提"],
        [
            ("内网地址", "是", "同 VPC 或通过终端节点打通的内部调用", "用户 VPC 终端节点 IP"),
            ("外网地址", "否", "本地、跨网或公网业务系统调用", "创建集群时开启公网入口并绑定 EIP"),
            ("内网域名", "否", "在 VPC 内用稳定域名代替内网 IP", "绑定自定义内网域名并与内网地址关联"),
            ("公网域名", "否", "互联网调用和可管理的公共入口", "已注册域名 + 已开启公网入口"),
        ],
        [1500, 1000, 3760, 3100],
    )
    add_bullet(doc, bullet_id, "如果从公网或其他安全组调用，入方向需按需放行 80/443。", "安全组：")
    add_bullet(doc, bullet_id, "如果后端服务在公网或其他安全组，出方向需放行后端地址与监听端口。", "后端连通：")
    add_bullet(doc, bullet_id, "物理机纳管网段与集群网段不一致时，需在集群详情中配置路由。", "路由：")
    add_bullet(doc, bullet_id, "采购时可选单 AZ 或 2-10 个 AZ 的多 AZ 方式；官网推荐多 AZ 以提升容灾能力。", "可用性：")
    add_source_body(doc, "专享集群创建完成后不支持切换 VPC、子网和安全组，因此这些是上线前必须确定的架构决策。", "S03")

    doc.add_heading("2.2 容量与运行规格", level=2)
    add_table(
        doc,
        ["专享集群规格", "节点", "最大支持发布API数", "官网标注延时"],
        [
            ("小规格 8U16G", "2", "500", "<20 ms"),
            ("中规格 16U32G", "2", "1000", "<15 ms"),
            ("大规格 32U64G", "2", "2000", "<10 ms"),
        ],
        [2700, 1100, 2850, 2710],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_source_body(doc, "上表是官网规格页给出的产品级参数；实际端到端响应还受网络、数据源、SQL复杂度和返回量影响，不应把该数值当作所有业务的应用级 SLA。", "S02")

    doc.add_heading("3 核心对象模型", level=1)
    add_callout(doc, "架构归纳", "数据服务的真正管理单元不是一个 API URL，而是“工作空间 + 专享集群 + API产品 + 消费应用 + 授权/审批 + 运行策略”共同组成的服务上下文。", fill=LIGHT_BLUE)
    object_rows = [
        ("DataArts Studio实例", "顶层产品载体", "工作空间、专享集群增量包", "提供数据治理的统一入口"),
        ("工作空间", "组织、权限和配额边界", "API、数据连接、审核人、集群共享关系", "确定谁可见、谁可开发与审核"),
        ("专享集群", "API运行时", "VPC、子网、安全组、ELB、节点、域名、日志/监控", "承载已发布 API 并提供稳定入口"),
        ("API目录", "API分类与检索单元", "API、服务目录路径", "组织开发和消费视图"),
        ("API", "数据服务产品", "路径、协议、请求方式、参数、取数逻辑、认证、版本、可见性", "定义对外稳定契约"),
        ("API编排工作流", "组合型 API", "入口 API、普通 API、条件、并行、输出处理算子", "无需复杂代码即可复用和重组服务"),
        ("应用 APP/IAM", "API消费身份", "AppKey/AppSecret 或 IAM 身份、已绑定 API", "将授权与真实调用方绑定"),
        ("授权", "API与消费身份的时效关系", "API、应用/白名单、截止时间、集群、Static参数", "决定谁在什么范围内可以调用"),
        ("审核单", "治理工作流", "发布、授权、续约、解除授权等操作", "把高风险变更从直接操作变为可审计流程"),
        ("流控/监控/日志", "运行保障对象", "API、用户、应用、集群、指标、访问记录", "保护后端、定位异常、支持审计"),
    ]
    add_table(doc, ["对象", "定位", "关联要素", "架构价值"], object_rows, [1700, 1900, 3330, 2430])

    doc.add_heading("3.1 关系主链", level=2)
    chains = [
        ("开发关系", "工作空间 → API目录 → API/编排 → 调试 → 版本 → 集群发布"),
        ("消费关系", "服务目录中的已发布API → 应用 → 授权/审批 → 认证信息 → 实际调用"),
        ("网络关系", "用户 VPC 终端节点或 EIP/域名 → ELB → 专享集群节点 → 后端数据源"),
        ("治理关系", "开发/调用操作 → 审核单 → 审批生效 → 监控/日志 → 变更或故障处理"),
    ]
    for label, text in chains:
        add_bullet(doc, bullet_id, text, label + "：")

    doc.add_heading("4 API提供方：服务生产与生命周期", level=1)
    add_source_body(doc, "官网给出的提供方主流程是：管理专享集群 → 新建审核人 → 创建 API → 调试 → 发布 → 管理/编排 → 可选流控 → 授权。", "S01")
    provider_steps = [
        ("准备运行时", "购买专享集群，决定 AZ、VPC、子网、安全组、公网入口，为工作空间分配 API 配额。", "集群可运行，网络可达"),
        ("准备数据连接", "在管理中心配置数据源连接，让 API 开发时能读取表与字段。", "数据源可连通"),
        ("组织 API", "建立 API 目录，设置名称、请求Path、协议、方法、认证、可见性、审核人、入参。", "服务契约成形"),
        ("实现取数逻辑", "选择配置式或脚本/MyBatis式，定义请求参数、返回参数、排序和分页。", "数据契约可执行"),
        ("调试", "在指定集群上填写参数并查看请求与返回结果；触发数据源与超时问题时回到开发页修正。", "功能与性能基线通过"),
        ("审核与发布", "选择已调试通过的集群发布；非审核人提交审批，审批生效后 API 才能被调用。", "API进入运行态"),
        ("运营与变更", "管理版本、可见性、停用/恢复、下线/删除、复制、同步、导入导出、授权、流控和监控。", "服务可持续演进"),
    ]
    add_table(doc, ["阶段", "主要动作", "出口条件"], provider_steps, [1650, 5360, 2350])

    doc.add_heading("4.1 两种主要API生产模式", level=2)
    add_table(
        doc,
        ["模式", "实现方式", "优势", "主要约束/适用性"],
        [
            ("配置式 API", "在界面中选数据源、连接、库表和字段，配置请求/返回/排序参数与操作符", "无需编码，上手快，输入和返回契约可视化", "适合单表或规则化查询；官网说明暂不支持 Hive 中文表/列场景"),
            ("脚本/MyBatis API", "编写 SQL 或 MyBatis 动态语句，再解析参数、绑定类型并测试 SQL", "可表达复杂逻辑，MyBatis 支持更丰富的 Body 参数类型", "需要 SQL 能力，必须审视动态条件、超时与返回量"),
            ("API编排", "拖拽入口 API、普通 API、条件、并行、输出处理算子形成工作流", "减少上层多次调用，支持结果映射和格式转换", "仅 3.0.6+ 专享集群；工作流无环、无孤立算子、最多20层"),
        ],
        [1650, 3100, 2250, 2360],
    )
    add_source_body(doc, "当前国际站用户指南的创建章节主要详解配置式和脚本/MyBatis式；API 参考的创建接口还枚举了注册、Groovy和编排类型。这里将用户指南中的两种交互式开发作为应用架构主干。", "S04/S05/S09/S19")
    add_figure(doc, inventory, "dataartsstudio_01_7696", 1, "图 6  API基本配置：名称、目录、路径、协议、认证、可见性与入参")
    add_figure(doc, inventory, "dataartsstudio_01_7696", 3, "图 7  从数据表字段选择请求、返回与排序参数")
    add_figure(doc, inventory, "dataartsstudio_01_7696", 8, "图 8  API在线测试：左侧配参，右侧查看请求与返回结果")

    doc.add_heading("4.2 API契约并非只是一个URL", level=2)
    contract_rows = [
        ("识别与路由", "API名称、API目录、请求Path、HTTP/HTTPS、GET/POST", "决定如何被发现和访问"),
        ("输入契约", "Query/Header/Path/Body/Static，类型、必填、默认值、空值、校验规则", "在进入后端前统一校验"),
        ("数据逻辑", "数据源/库表/字段/操作符，或 SQL/MyBatis 脚本", "将应用逻辑与数据访问逻辑解耦"),
        ("输出契约", "返回字段名、类型、排序、分页、是否返回总条数", "稳定上层应用依赖"),
        ("治理属性", "审核人、安全认证、可见性、标签、访问日志、最低保留期限", "把发布、共享、审计和下线放进同一契约"),
    ]
    add_table(doc, ["契约维度", "包含内容", "架构作用"], contract_rows, [1750, 4920, 2690])
    add_source_body(doc, "请求参数总大小也有明确限制：Query+Path 最大 32KB，Header 最大 128KB，Body 最大 128KB。", "S04")

    doc.add_heading("4.3 调试、审核、发布与版本演进", level=2)
    add_source_body(doc, "API 只有发布后才能被调用。发布时选择已调试通过的集群；若操作人不具备审核人权限，发布会进入审批。", "S06/S07/S17")
    add_figure(doc, inventory, "dataartsstudio_01_0308", 1, "图 9  发布API时选择调试通过的专享集群")
    add_source_body(doc, "专享版通过“编辑已发布API→再次调试→再发布”形成新版本。系统最多保留最近10条版本记录，支持查看、删除、发布和两版本对比。", "S08")
    add_figure(doc, inventory, "dataartsstudio_01_0339", 3, "图 10  API版本管理：历史版本、发布与版本对比入口")

    doc.add_heading("4.4 API编排是平台内的轻量服务组合层", level=2)
    add_source_body(doc, "API 编排通过拖拽工作流把多个已开发 API 以串行或并行方式组合，并通过入口 API 对外提供一个统一结果。", "S09")
    add_table(
        doc,
        ["算子", "是否必须", "位置/数量约束", "作用"],
        [
            ("入口API", "必须", "只能1个，位于最上游，向下1个分支", "定义工作流名称、URL、认证与请求参数"),
            ("普通API", "至少1个", "位于中间层，上下游均有算子", "执行数据查询并传递请求/结果变量"),
            ("条件分支", "可选", "2-20个分支，多分支满足时只执行第一个", "基于上游参数或结果集选择路径"),
            ("并行处理", "可选", "2-20个分支，必须配置失败策略", "并发执行互不影响的多分支逻辑"),
            ("输出处理", "必须", "只能1个，位于最下游，直接上游必须是普通API", "错误码映射、结果集映射和格式转换"),
        ],
        [1600, 1200, 3300, 3260],
    )
    add_figure(doc, inventory, "dataartsstudio_01_0324", 1, "图 11  API工作流编排页：算子面板、画布与算子配置面板")

    doc.add_heading("5 API调用方：服务发现、授权与安全", level=1)
    add_source_body(doc, "调用方从“调用API > 服务目录”发现已发布API。如果开发者没有主动授权，调用方需将API申请绑定到符合认证类型的应用，待审批后才可调用。", "S12")

    doc.add_heading("5.1 认证方式对比", level=2)
    add_table(
        doc,
        ["认证方式", "安全级别", "授权/认证机制", "调用方式", "建议"],
        [
            ("APP认证", "高", "API授权给APP应用，使用AppKey/AppSecret签名", "推荐SDK；也可使用API工具生成签名后调用", "正式系统首选"),
            ("IAM认证", "中", "API授权给IAM应用或白名单，通过IAM服务获取用户Token", "API工具携带Token调用", "适合账号身份集成"),
            ("无认证", "低", "不需授权，任意用户可访问", "API工具；Query/Path入参时可用浏览器", "仅测试，不推荐正式使用"),
        ],
        [1400, 1100, 3110, 2400, 1350],
    )
    add_source_body(doc, "官网明确推荐“APP认证 + SDK调用”。无认证存在数据泄露、数据库高并发和 SQL 注入等风险，不应作为正式业务的常规入口。", "S10")

    doc.add_heading("5.2 应用是运行时授权主体", level=2)
    add_body(doc, "应用并非简单的分组标签，它承载 API 授权和调用凭证。APP 认证API只能授权给APP类型应用，IAM认证API只能授权给IAM类型应用，也可按白名单授权 IAM 调用方。")
    add_figure(doc, inventory, "dataartsstudio_01_0333", 4, "图 12  将APP认证API授权给应用：截止时间、集群、Static参数与应用选择")
    add_body(doc, "授权关系还可包含截止时间、发布集群以及 Static 入参值。因此“API可见”不等于“API可调用”：可见性控制服务目录发现，授权与认证控制实际访问。")

    doc.add_heading("5.3 推荐调用路径", level=2)
    call_steps = [
        "在服务目录选择已发布API，核对路径、方法、入参和可见性。",
        "创建或选择与API认证类型一致的应用，申请或由开发者主动授权。",
        "审核通过后获取调用地址、AppKey/AppSecret或IAM Token，同时确认授权截止时间。",
        "使用官网推荐SDK签名调用APP认证API，对超时、重试、状态码和分页做应用端处理。",
        "根据访问日志、调用趋势、延时和4xx/5xx指标持续运营授权与调用。",
    ]
    for step in call_steps:
        add_numbered(doc, decimal_id, step)
    add_figure(doc, inventory, "dataartsstudio_01_7697", 3, "图 13  从API调用信息页下载SDK代码示例")

    doc.add_heading("6 治理、可观测与运营架构", level=1)
    doc.add_heading("6.1 审核中心：把发布和授权放进可审计流程", level=2)
    add_source_body(doc, "审核中心面向 API 开放方与调用方。当发布人没有审核人权限时，发布要提交审核；具备审核人权限时可直接发布。发起人可在待审核阶段撤销申请。", "S17")
    add_figure(doc, inventory, "dataartsstudio_01_0312", 2, "图 14  审核中心待审核列表与“审核”入口")
    add_callout(doc, "角色边界", "审核人可来自工作空间管理员或开发者；运维者和访客不能添加为审核人。工作空间管理员默认也没有审批权，需将自己加入审核人。", fill=CALLOUT)

    doc.add_heading("6.2 流控策略：独立于API的后端保护对象", level=2)
    add_source_body(doc, "流控策略可按用户、应用和时间段维度限制API调用次数。策略与API相互独立，只有绑定后才生效；同一环境中一个API只能绑定一个流控策略，一个策略可绑定多个API。", "S13")
    add_table(
        doc,
        ["限流维度", "含义", "关系"],
        [
            ("API流量限制", "单个API在单位时间内的请求上限", "顶层总量边界"),
            ("用户流量限制", "单个用户在单位时间内的请求上限", "不超过API流量限制"),
            ("应用流量限制", "单个APP在单位时间内的请求上限", "不超过用户流量限制"),
        ],
        [2200, 4460, 2700],
    )
    add_figure(doc, inventory, "dataartsstudio_01_0311", 1, "图 15  创建API流控策略：时长、API、用户与应用维度限制")

    doc.add_heading("6.3 两级可观测：产品运营 + 集群资源", level=2)
    add_table(
        doc,
        ["观测层", "入口", "关键内容", "用途"],
        [
            ("API/APP运营层", "数据服务总览、访问日志", "发布数、授权应用、总调用、成功/失败、合法/非法、TOP5、单次请求详情", "理解服务被谁、以什么结果消费"),
            ("集群运行层", "Cloud Eye（CES）", "平均/最大延时、调用次数、4xx/5xx、健康检查、CPU、内存、网络连接", "理解专享集群是否健康和需要扩容"),
            ("日志留存层", "集群节点、LTS或OBS转储", "API访问记录、请求ID、API/应用、状态码、请求时长", "故障定位、调用追溯和长期审计"),
        ],
        [1800, 1700, 3980, 1880],
    )
    add_source_body(doc, "Cloud Eye 监控指标命名空间为 SYS.DLM，原始指标周期为1分钟。需先在专享集群中开启 CES 监控。", "S14/S15")
    add_figure(doc, inventory, "dataartsstudio_01_7692", 1, "图 16  Cloud Eye中的数据服务集群监控指标")
    add_source_body(doc, "API 访问日志默认保存在集群节点，专享版可转储到 LTS 或 OBS。产品页访问日志仅保留7天，长期审计应使用转储架构。", "S16")
    add_figure(doc, inventory, "dataartsstudio_01_0900", 3, "图 17  LTS中的API访问日志详情")

    doc.add_heading("6.4 可见性、授权与认证是三个不同控制点", level=2)
    add_table(
        doc,
        ["控制点", "回答的问题", "主要选项/对象", "不能替代"],
        [
            ("API可见性", "谁能在服务目录看到API", "当前工作空间、当前项目、当前租户；可额外按项目ID设置", "不等于已授权"),
            ("API授权", "某个应用/账号是否允许调用", "APP应用、IAM应用、IAM白名单、授权截止时间", "不等于调用请求已认证"),
            ("请求认证", "这次请求是否属于已授权主体", "AppKey/AppSecret签名、IAM Token、无认证", "不等于业务参数合法"),
            ("参数校验/流控", "请求内容与调用频率是否可接受", "类型、长度/枚举/正则、必填、API/用户/应用限额", "不等于后端数据一定有结果"),
        ],
        [1750, 2590, 3330, 1690],
    )
    add_source_body(doc, "上表的四点区分是架构归纳，官网分别在创建API、设置可见、授权、调用方式和流控页面描述了各自机制。", "S04/S10/S11/S13/S20")

    doc.add_heading("7 开放接口与集成架构", level=1)
    add_source_body(doc, "除控制台外，DataArts Studio 还开放数据服务管理 API，便于用代码管理 API、应用、目录、授权和集群。这些是控制面接口，与业务应用调用已发布数据API的运行时入口不同。", "S18")
    api_groups = [
        ("API管理接口", "创建、查询、更新、发布、调试与生命周期管理"),
        ("申请管理接口", "处理API申请与审批关系"),
        ("消息管理接口", "处理消息和通知"),
        ("授权管理接口", "查询API已授权APP及APP已拥有API"),
        ("服务目录管理接口", "创建、更新、查询、移动、删除目录与API"),
        ("网关管理接口", "管理数据服务网关相关信息"),
        ("应用管理接口", "创建、查询与管理调用应用"),
        ("总览接口", "获取数据服务总览统计"),
        ("集群管理接口", "查询和管理专享集群相关信息"),
    ]
    add_table(doc, ["官方接口类型", "在应用架构中的作用"], api_groups, [3000, 6360])
    add_source_body(doc, "创建API的控制面路径为 POST /v1/{project_id}/service/apis，请求头包含 X-Auth-Token 和必填的 workspace ID；Dlm-Type 可选用于指定 SHARED 或 EXCLUSIVE。", "S19")
    add_callout(doc, "集成建议", "将控制面自动化与业务数据API调用分开：前者由平台工程和发布流水线使用，后者由业务应用在授权、认证和流控边界内调用。", fill=LIGHT_BLUE)

    doc.add_heading("8 非功能边界与选型判断", level=1)
    doc.add_heading("8.1 返回量与分页边界", level=2)
    add_table(
        doc,
        ["API类型", "场景", "数据源", "官网默认/自定义规格"],
        [
            ("配置类", "调试", "DLI/DWS/MySQL/RDS/Hive/HBase/ClickHouse/Hetu/GBase/Doris/Oracle", "10条"),
            ("配置类", "调用", "同上", "100条"),
            ("脚本/MyBatis类", "测试SQL", "-", "10条"),
            ("脚本/MyBatis类", "调试/调用", "DLI", "默认分页100；自定义分页1000"),
            ("脚本/MyBatis类", "调试/调用", "DWS/MySQL/RDS/Hive/HBase/ClickHouse/Hetu/GBase/Doris/Oracle", "默认分页10；自定义分页2000"),
        ],
        [1850, 1300, 3900, 2310],
    )
    add_source_body(doc, "这些限制再次说明数据服务是交互式API，不是大规模数据导出通道。", "S02")

    doc.add_heading("8.2 关键架构质量属性", level=2)
    quality_rows = [
        ("可用性", "集群多节点 + ELB；采购时可选多 AZ", "通过监控和容量测试验证后端数据源是否也具备对等高可用"),
        ("性能", "官网给出集群规格、API数与延时标注；Cloud Eye有平均/最大延时、CPU、内存等指标", "用业务SQL、真实返回量和并发建立性能基线"),
        ("安全", "VPC/安全组/HTTPS + APP/IAM + 可见性/授权/审批 + 参数校验", "正式环境禁用无认证，缩小入方向范围，定期轮换应用密钥"),
        ("可运维性", "版本、编辑再发布、停用/下线、导入导出、访问日志与集群监控", "把API发布和回退放入变更流程，并对版本与数据契约做回归测试"),
        ("审计性", "审核中心 + API/APP视角运营 + 访问日志转储", "根据合规保留期设计 LTS/OBS 长期存储，不依赖产品页7天保留"),
        ("治理性", "工作空间、API目录、可见范围、应用、授权、标签、最低保留期限", "明确API责任人、审核人、消费应用、数据口径与下线通知责任"),
    ]
    add_table(doc, ["质量属性", "产品机制", "落地要点"], quality_rows, [1600, 4300, 3460])

    doc.add_heading("9 建议的企业落地蓝图", level=1)
    add_callout(doc, "建议架构", "生产环境采用专享版多 AZ 集群，以私网调用为主；只为确有跨网需求的API开启公网入口。正式调用默认APP认证+SDK，API发布和授权均经审批，并同时开启Cloud Eye、LTS/OBS转储与流控。", fill=LIGHT_BLUE, accent=DARK_BLUE)
    blueprint_steps = [
        ("阶段1：建边界", "定义工作空间、产品/项目责任、IAM角色、审核人、API目录和命名规则。"),
        ("阶段2：建运行时", "选择集群规格与 AZ，确定 VPC/子网/安全组，预留端点、域名和后端连通。"),
        ("阶段3：建API产品模板", "统一 HTTPS、APP认证、参数校验、分页、错误码、日志、保留期和可见性默认值。"),
        ("阶段4：建发布门禁", "把在线调试、数据契约检查、审批、版本记录和回退预案变成上线必经步骤。"),
        ("阶段5：建消费运营", "将应用与真实消费方绑定，设置授权截止时间、流控和密钥轮换，按APP视角复盘消费。"),
        ("阶段6：建可观测闭环", "以 Cloud Eye 指标观测集群，以总览观测API/APP，以LTS/OBS保留访问证据，为延时和4xx/5xx建立告警。"),
    ]
    for title_text, body in blueprint_steps:
        add_numbered(doc, decimal_id, body, title_text + " ")

    doc.add_heading("9.1 上线前架构检查清单", level=2)
    checklist = [
        "VPC、子网、安全组和 AZ 已确定，因为集群创建后不支持切换。",
        "从调用方到API入口、再到后端数据源的入/出方向网络均已验证。",
        "API目录、责任人、审核人、可见性、最低保留期限和下线通知流程已定义。",
        "正式API使用HTTPS与APP/IAM认证，无认证仅用于受控测试。",
        "请求参数类型、必填、校验、默认值与返回契约已回归，不只验证正常路径。",
        "目标并发、后端数据源容量、返回量与分页已通过真实SQL负载测试。",
        "API/用户/应用三级流控值已与后端实际承载能力匹配。",
        "Cloud Eye、延时与4xx/5xx告警、LTS/OBS转储和日志保留期已开启并验证。",
        "版本变更、下线/停用、授权解除、密钥轮换和故障回退已有可执行预案。",
    ]
    for item in checklist:
        add_bullet(doc, bullet_id, item)

    doc.add_heading("9.2 不建议使用数据服务的场景", level=2)
    for label, text in [
        ("大批量数据导出：", "官网定位是小批量快速交互，大量数据应使用共享交换、文件分发或其他批处理方案。"),
        ("对后端无容量约束的公开查询：", "无认证与无流控会把数据泄露、SQL注入和数据库高并发风险直接暴露给上层。"),
        ("需要复杂长时编排的业务流：", "产品的 API 编排适合轻量串并行和映射转换，不应默认代替长事务、耗时任务与复杂补偿引擎。"),
    ]:
        add_bullet(doc, bullet_id, text, label)

    doc.add_heading("10 结论", level=1)
    add_body(doc, "DataArts Studio 数据服务的中心不是“生成一个查询接口”，而是把数据连接、服务契约、运行集群、消费身份与治理证据收拢到同一个生命周期。")
    add_body(doc, "从应用架构角度，最关键的四个设计点是：把专享集群视为独立运行时；把 API 视为带治理属性的数据产品；把应用与授权视为消费边界；把审批、流控、监控和日志视为上线必需的运行闭环。")
    add_callout(doc, "最终判断", "如果目标是用可管理、可授权、可审计的方式向上层应用开放小批量数据查询，这套架构与目标高度匹配；如果目标是大量数据搬运、不受控的匿名开放或长时复杂流程编排，则应设计其他专用通道。", fill=LIGHT_BLUE, accent=DARK_BLUE)

    doc.add_page_break()
    doc.add_heading("附录A  官网“数据服务”页面地图", level=1)
    add_body(doc, f"本次从“数据服务”根页面递归读取 {inventory['page_count']} 个官方页面，共收集 {inventory['unique_image_count']} 张唯一官网原图。下表用于回溯功能范围，不表示所有页面都在正文逐张展示。")
    page_map_rows = []
    for index, page in enumerate(inventory["pages"], start=1):
        depth_label = {0: "根页", 1: "一级", 2: "二级", 3: "三级"}.get(page["depth"], f"{page['depth']}级")
        page_map_rows.append(
            (
                str(index),
                depth_label,
                page["title"],
                str(len(page["images"])),
                page.get("updated") or "-",
            )
        )
    add_table(
        doc,
        ["序号", "层级", "页面标题", "原图数", "官网更新时间"],
        page_map_rows,
        [650, 900, 4400, 900, 2510],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )

    doc.add_page_break()
    doc.add_heading("附录B  主要官方来源", level=1)
    add_body(doc, "以下链接为正文架构判断的主要证据页。截图均下载自对应页面内的原始产品图，未使用重绘界面冒充产品截图。")
    add_source_list(doc, SOURCES)

    doc.save(OUTPUT)
    print(f"output={OUTPUT}")
    print(f"pages_researched={inventory['page_count']}")
    print(f"official_images_collected={inventory['unique_image_count']}")
    print("official_images_embedded=16")


if __name__ == "__main__":
    build_document()
