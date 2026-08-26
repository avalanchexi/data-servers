from __future__ import annotations

import json
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_volcengine_data_service_report import (
    BLUE,
    GOLD,
    INK,
    MUTED,
    SOURCES,
    add_caption,
    add_hyperlink,
    set_font,
    set_picture_alt,
)


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "火山引擎DataLeap数据服务业务架构梳理.docx"
OUTPUT = ROOT / "火山引擎DataLeap数据服务业务架构梳理（含页面截图）.docx"
IMAGE_DIR = ROOT / ".tmp" / "volcengine-data-service-report" / "page-screenshots"
INVENTORY = IMAGE_DIR / "inventory.json"

# One representative product-interface image per page. Values are 1-based image
# positions in the official document body. The original bytes are embedded as-is.
SELECTIONS = {
    "S3": 10,
    "S4": 2,
    "S5": 2,
    "S6": 1,
    "S7": 2,
    "S8": 3,
    "S10": 2,
    "S11": 1,
    "S12": 3,
    "S14": 2,
    "S15": 1,
    "S16": 1,
    "S17": 1,
    "S18": 1,
    "S19": 1,
    "S22": 2,
    "S23": 1,
}

SCREENSHOT_DESCRIPTIONS = {
    "S3": "API 详情中的调用说明、环境地址与代码示例",
    "S4": "数据源列表及新增数据源入口",
    "S5": "数据源列表中的切流操作入口",
    "S6": "物理表创建及待选表配置界面",
    "S7": "逻辑表详情与字段信息界面",
    "S8": "API 开发中的逻辑表与请求参数配置界面",
    "S10": "API 调用说明、环境地址和示例代码界面",
    "S11": "API 开发页中的运维入口",
    "S12": "OneService SQL 编辑器中的数组参数配置示例",
    "S14": "API 集市的资产检索与列表界面",
    "S15": "逻辑表集市的资产检索与列表界面",
    "S16": "账户权限管理列表界面",
    "S17": "项目管理列表及项目操作界面",
    "S18": "业务线管理列表界面",
    "S19": "应用管理列表及密钥、授权相关入口",
    "S22": "审批中心的待审批工单列表",
    "S23": "标签管理列表及新建标签组入口",
}

NO_SCREENSHOT_NOTES = {
    "S1": "该页为产品定位与价值说明，官方正文未提供产品界面截图。",
    "S2": "该页为核心概念说明，官方正文未提供产品界面截图。",
    "S9": "该页为 API 编排开发说明，当前官方正文未嵌入产品界面截图。",
    "S13": "该页为 Dynamic SQL 语法说明，官方正文未提供产品界面截图。",
    "S20": "该页为公网配置说明，当前官方正文未嵌入产品界面截图。",
    "S21": "该页为 VPC 配置说明，当前官方正文未嵌入产品界面截图。",
}


def set_no_picture_compression(doc: Document) -> None:
    """Tell Word to retain full-resolution image data when the file is edited."""
    settings = doc.settings.element
    node = settings.find(qn("w:doNotCompressPictures"))
    if node is None:
        node = OxmlElement("w:doNotCompressPictures")
        settings.append(node)
    node.set(qn("w:val"), "true")


def add_source_line(doc: Document, sid: str, page_id: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("官方页面：")
    set_font(run, size=9.2, bold=True, color=INK)
    url = f"https://docs.volcengine.com/docs/6260/{page_id}?lang=zh"
    add_hyperlink(paragraph, f"{sid} · {url}", url, size=9.2)


def fit_size(width_px: int, height_px: int) -> tuple[float, float]:
    max_width = 6.25
    max_height = 6.05
    aspect = width_px / height_px
    width = max_width
    height = width / aspect
    if height > max_height:
        height = max_height
        width = height * aspect
    return width, height


def add_screenshot(doc: Document, sid: str, name: str, image_info: dict, figure_no: int) -> None:
    image_path = IMAGE_DIR / image_info["localName"]
    with Image.open(image_path) as source:
        width_px, height_px = source.size
    display_width, display_height = fit_size(width_px, height_px)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(
        str(image_path),
        width=Inches(display_width),
        height=Inches(display_height),
    )
    description = SCREENSHOT_DESCRIPTIONS[sid]
    set_picture_alt(
        paragraph,
        f"[{sid}] {name}：{description}。火山引擎官方文档产品界面截图，"
        f"原始分辨率 {width_px}×{height_px} 像素，未降采样。",
    )
    add_caption(
        doc,
        f"图 C-{figure_no} [{sid}] {name}：{description}（火山引擎官方文档原图；"
        f"{width_px}×{height_px} px；DOCX 内按原始文件直接嵌入、未降采样）",
    )


def add_no_screenshot_note(doc: Document, note: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("截图状态  ")
    set_font(run, size=10, bold=True, color=GOLD)
    run = paragraph.add_run(note)
    set_font(run, size=10, color=INK)


def main() -> None:
    inventory = json.loads(INVENTORY.read_text(encoding="utf-8"))
    page_lookup = {page["sid"]: page for page in inventory["pages"]}

    for sid, selected_position in SELECTIONS.items():
        images = page_lookup[sid]["images"]
        if selected_position < 1 or selected_position > len(images):
            raise ValueError(f"Invalid image selection for {sid}: {selected_position}")
        image_path = IMAGE_DIR / images[selected_position - 1]["localName"]
        if not image_path.is_file():
            raise FileNotFoundError(image_path)

    covered = set(SELECTIONS) | set(NO_SCREENSHOT_NOTES)
    expected = {source[0] for source in SOURCES}
    if covered != expected:
        raise ValueError(f"Screenshot coverage mismatch: missing={expected - covered}, extra={covered - expected}")

    doc = Document(INPUT)
    set_no_picture_compression(doc)
    doc.add_page_break()
    doc.add_heading("附录 C 各页面官方产品截图", level=1)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(
        "本附录逐页核对“数据服务”目录。23 个页面中，17 页的官方正文提供了产品界面原图，"
        "每页选取 1 张最能代表该页面业务动作的截图；其余 6 页未提供界面图，已保留页面条目并说明。"
    )
    set_font(run, size=10.5, color=INK)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(
        "清晰度说明：所有截图均从火山引擎官方文档的原始图片地址下载，并以原始文件字节直接嵌入 DOCX；"
        "仅设置版面显示宽高，不执行降采样、转码或有损压缩，且已启用“不压缩文件中的图像”设置。"
    )
    set_font(run, size=9.6, bold=True, color=BLUE)

    current_module = None
    figure_no = 0
    first_entry = True
    for sid, name, module, page_id, _ in SOURCES:
        if module != current_module:
            if not first_entry:
                doc.add_page_break()
            doc.add_heading(module, level=2)
            current_module = module
        elif not first_entry:
            doc.add_page_break()

        doc.add_heading(f"{sid} {name}", level=3)
        add_source_line(doc, sid, page_id)

        if sid in SELECTIONS:
            figure_no += 1
            selected = page_lookup[sid]["images"][SELECTIONS[sid] - 1]
            add_screenshot(doc, sid, name, selected, figure_no)
        else:
            add_no_screenshot_note(doc, NO_SCREENSHOT_NOTES[sid])
        first_entry = False

    doc.core_properties.title = "火山引擎 DataLeap 数据服务业务架构梳理（含页面截图）"
    doc.core_properties.subject = "基于火山引擎官方文档的数据服务业务架构分析及页面原图附录"
    doc.core_properties.keywords = "火山引擎, DataLeap, 数据服务, 业务架构, 官方截图, 原图"
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Embedded screenshots: {figure_no}")


if __name__ == "__main__":
    main()
